"""Build CV Module Expert Review Word document."""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

OUTPUT = os.path.join(
    os.path.dirname(__file__),
    "CV_Module_Expert_Review_2026-02-14.docx"
)

def set_cell_shading(cell, color_hex):
    """Set background color for a table cell."""
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex,
        qn('w:val'): 'clear',
    })
    shading.append(shading_elem)


def add_table_row(table, cells_text, bold=False, header=False):
    """Add a row to a table."""
    row = table.add_row()
    for i, text in enumerate(cells_text):
        cell = row.cells[i]
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        run.font.size = Pt(9)
        if bold or header:
            run.bold = True
        if header:
            # Light blue header
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            shading = tcPr.makeelement(qn('w:shd'), {
                qn('w:fill'): '4472C4',
                qn('w:val'): 'clear',
            })
            tcPr.append(shading)
            run.font.color.rgb = RGBColor(255, 255, 255)
    return row


def build_document():
    doc = Document()

    # -- Page margins --
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # -- Styles --
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.15

    for level in range(1, 4):
        hstyle = doc.styles[f'Heading {level}']
        hstyle.font.name = 'Calibri'
        if level == 1:
            hstyle.font.size = Pt(16)
            hstyle.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        elif level == 2:
            hstyle.font.size = Pt(13)
            hstyle.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
        else:
            hstyle.font.size = Pt(11)
            hstyle.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

    # ============================================================
    # TITLE
    # ============================================================
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run('LLM-ABM 構念與驗證模組 (C&V Module)')
    run.font.size = Pt(20)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle_p.add_run('專家審查綜合報告')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run('2026-02-14')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.add_paragraph()  # spacer

    # ============================================================
    # 1. 審查概述
    # ============================================================
    doc.add_heading('1. 審查概述', level=1)

    items = [
        ('日期', '2026-02-14'),
        ('模組', 'compute_validation_metrics.py (~1,297 行)'),
        ('審查團隊', 'LLM 行為工程專家、行為/社會科學專家、水資源工程專家、資工系教授'),
        ('目標', '通用性、可擴充性 (10K+ agents)、多理論支援、具體範例'),
    ]
    for label, value in items:
        p = doc.add_paragraph()
        run = p.add_run(f'{label}：')
        run.bold = True
        p.add_run(value)

    p = doc.add_paragraph()
    p.add_run(
        '本報告綜合四位專家對 C&V Module 的獨立審查意見，'
        '涵蓋程式架構、行為科學理論、水資源領域特性、及 LLM 特有驗證需求。'
        '所有建議依 priority (P0-P3) 分級，其中 P0 為必須立即修復的 blocking bug。'
    )

    # ============================================================
    # 2. 跨專家共識 (P0)
    # ============================================================
    doc.add_heading('2. 跨專家共識 (P0 — 必須立即修復)', level=1)

    doc.add_heading('2.1 EBE 平均計算 Bug', level=2)
    p = doc.add_paragraph()
    p.add_run('問題：').bold = True
    p.add_run(
        '目前程式碼以 (EBE_owner + EBE_renter) / 2 計算整體 Empirical Behavioral Entropy (EBE)，'
        '但 Shannon entropy 不具有可加性 (non-additive)。兩個子群體的 entropy 算術平均不等於合併分佈的 entropy。'
    )
    p = doc.add_paragraph()
    p.add_run('修正方案：').bold = True
    p.add_run(
        '應先將 owner 與 renter 的動作分佈依 population ratio 加權合併，'
        '再對合併後的分佈重新計算 Shannon entropy。'
        '公式：H(X) = -Σ p_merged(x) log₂ p_merged(x)，其中 p_merged = w_owner · p_owner + w_renter · p_renter。'
    )

    doc.add_heading('2.2 TP/CP 提取失敗 Silent Default', level=2)
    p = doc.add_paragraph()
    p.add_run('問題：').bold = True
    p.add_run(
        '當 LLM 回應中無法提取 Threat Perception (TP) 或 Coping Perception (CP) label 時，'
        '程式碼 silently defaults to "M" (Medium)。'
        '這導致 CACR (Construct-Action Coherence Rate) 被系統性地高估，'
        '因為 "M" 與大多數動作的 coherence mapping 吻合度高。'
    )
    p = doc.add_paragraph()
    p.add_run('修正方案：').bold = True
    p.add_run(
        '使用 "UNKNOWN" sentinel value。在 CACR 計算中排除所有 TP 或 CP = "UNKNOWN" 的 traces。'
        '另外報告 extract_failure_rate 作為獨立的 data quality metric。'
    )

    doc.add_heading('2.3 Agent Type 推斷循環性', level=2)
    p = doc.add_paragraph()
    p.add_run('問題：').bold = True
    p.add_run(
        '目前以 proposed action 推斷 agent type (owner vs. renter)，'
        '例如看到 "elevation" 就判斷為 owner。這是 circular reasoning：'
        '用行為推斷身份，再用身份驗證行為。'
    )
    p = doc.add_paragraph()
    p.add_run('修正方案：').bold = True
    p.add_run(
        '應從 agent profile (initialization data) 中 join agent_type 欄位。'
        '所有 agent type 資訊應在實驗開始時就確定，而非從 trace data 中推斷。'
    )

    # ============================================================
    # 3. 架構重構建議 (P1)
    # ============================================================
    doc.add_heading('3. 架構重構建議 (P1)', level=1)

    doc.add_heading('3.1 模組化拆分 (CS 教授)', level=2)
    p = doc.add_paragraph()
    p.add_run(
        '目前 compute_validation_metrics.py 為 1,297 行的 monolith，'
        '違反 Single Responsibility Principle。建議拆分為 broker/validators/validation/ package，'
        '包含以下 7 個子模組：'
    )
    items = [
        'theories/ — 行為理論 coherence specs (PMT, TPB, HBM, PT, PADM)',
        'hallucinations/ — LLM hallucination & constraint violation detectors',
        'benchmarks/ — 各 empirical benchmark 獨立 callable',
        'metrics/ — 通用 metric 計算工具 (entropy, Moran\'s I, survival analysis)',
        'io/ — Trace 讀取、streaming、格式轉換',
        'reporting/ — 報表生成、Excel/Markdown 輸出',
        'engine.py — ValidationRunner facade，整合上述模組',
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    p = doc.add_paragraph()
    p.add_run('與現有 BenchmarkRegistry 整合：').bold = True
    p.add_run(
        '現有 empirical_benchmarks.py 中的 BenchmarkRegistry 已提供 benchmark 註冊機制，'
        '新架構應將其作為 benchmarks/ 子模組的基礎，避免重複建設。'
    )

    doc.add_heading('3.2 可插拔理論支援 (所有專家一致)', level=2)
    p = doc.add_paragraph()
    p.add_run(
        '定義 BehavioralTheory protocol，使不同行為理論以統一介面整合：'
    )
    # Code block
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    run = p.add_run(
        'class BehavioralTheory(Protocol):\n'
        '    def get_coherent_actions(\n'
        '        self, constructs: Dict[str, str], agent_type: str\n'
        '    ) -> List[str]: ...'
    )
    run.font.name = 'Consolas'
    run.font.size = Pt(9)

    p = doc.add_paragraph()
    p.add_run('支援的理論框架：').bold = True

    theories = [
        ('PMT (Protection Motivation Theory)', '2D lookup (TP × CP → coherent actions)'),
        ('TPB (Theory of Planned Behavior)', '3D lookup (Attitude × SubjectiveNorm × PBC)'),
        ('HBM (Health Belief Model)', '4-6D scoring (Susceptibility, Severity, Benefits, Barriers, Cues, Self-efficacy)'),
        ('PADM (Protective Action Decision Model)', 'Process-based (Risk identification → Assessment → Decision)'),
        ('PT (Prospect Theory)', 'Frame-conditional (Loss/Gain frame × Reference point → Risk preference)'),
    ]
    for name, desc in theories:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f'{name}：')
        run.bold = True
        p.add_run(desc)

    p = doc.add_paragraph()
    p.add_run('社科專家關鍵指出：').bold = True
    p.add_run(
        '行為理論分為兩種 paradigm — '
        'Construct-Action Mapping (PMT/TPB/HBM) 與 Frame-Conditional (PT/Nudge)。'
        '前者可用 lookup table 實作，後者需要 conditional logic。'
        '統一 protocol 必須同時支援兩種 paradigm。'
    )
    p = doc.add_paragraph()
    p.add_run(
        '建議使用 TheoryRegistry + YAML 宣告式配置，'
        '使用者只需提供 theory name 與 coherence table，不需修改 Python 程式碼。'
    )

    doc.add_heading('3.3 Benchmark 外部化', level=2)
    p = doc.add_paragraph()
    p.add_run(
        '目前 benchmark 計算內嵌於 ~410 行的 switch statement 中，'
        '難以維護且無法由使用者自訂。建議定義 BenchmarkComputation protocol：'
    )
    p = doc.add_paragraph()
    run = p.add_run(
        'class BenchmarkComputation(Protocol):\n'
        '    name: str\n'
        '    range: Tuple[float, float]\n'
        '    weight: float\n'
        '    def compute(self, traces, profiles) -> float: ...'
    )
    run.font.name = 'Consolas'
    run.font.size = Pt(9)

    p = doc.add_paragraph()
    p.add_run(
        '每個 benchmark 獨立為一個 callable plugin，並以 YAML 格式定義 '
        'benchmark name, range, weight, filter criteria, metric type。'
        '替換現有 switch statement，提升可維護性與擴展性。'
    )

    # ============================================================
    # 4. LLM 特定驗證缺口
    # ============================================================
    doc.add_heading('4. LLM 特定驗證缺口 (P1-P2, LLM 行為工程專家)', level=1)

    doc.add_heading('4.1 Sycophancy Detection', level=2)
    p = doc.add_paragraph()
    p.add_run(
        'LLM 可能僅 mirror prompt framing 而非真正推理。'
        '例如，若 prompt 強調 "flood risk is high"，LLM 可能直接輸出高風險適應行動，'
        '而非基於 agent state 獨立判斷。'
    )
    p = doc.add_paragraph()
    p.add_run('建議：').bold = True
    p.add_run(
        '加入 L1.5 Prompt-Response Independence (PRI) 指標。'
        '方法：對同一 agent state 使用不同 prompt framing (neutral vs. risk-emphasizing)，'
        '比較 action distribution 差異。PRI = 1 - JSD(P_neutral, P_framed)，'
        '其中 JSD 為 Jensen-Shannon Divergence。PRI > 0.7 為 acceptable。'
    )

    doc.add_heading('4.2 CACR Temporal Decomposition', level=2)
    p = doc.add_paragraph()
    p.add_run(
        '整體 CACR 可能掩蓋 temporal degradation。'
        '隨著模擬進行，LLM 可能因 context window 限制或 instruction-following fatigue '
        '而降低 construct-action coherence。'
    )
    p = doc.add_paragraph()
    p.add_run('建議：').bold = True
    p.add_run(
        '計算 CACR-by-year，繪製時序圖。'
        '若 CACR 斜率 < -0.02/year，標記為 "instruction-following collapse at scale" 風險。'
    )

    doc.add_heading('4.3 Construct Label Circularity', level=2)
    p = doc.add_paragraph()
    p.add_run('社科 + LLM 專家共識：').bold = True
    p.add_run(
        '目前 LLM 自己產生 TP/CP labels，又用這些 labels 驗證自己的 action coherence。'
        '這測量的是 self-consistency，而非 construct validity。'
    )
    p = doc.add_paragraph()
    p.add_run('建議：').bold = True
    p.add_run(
        '實施 "construct grounding" 驗證 — 使用外部 ground truth '
        '(如已知 flood depth → expected TP range) 驗證 LLM 產生的 construct labels 是否合理。'
        '例如：flood depth > 3ft 且 TP = "L" 應標記為 grounding failure。'
    )

    doc.add_heading('4.4 Reasoning Faithfulness', level=2)
    p = doc.add_paragraph()
    p.add_run(
        'LLM 可能輸出 TP_LABEL = "VH" (Very High) 但 reasoning 文字寫 '
        '"flood risk is minimal"。這種 label-reasoning disconnect 表示 LLM '
        '不是真正理解 construct，只是在 pattern match label format。'
    )
    p = doc.add_paragraph()
    p.add_run('建議：').bold = True
    p.add_run(
        '加入 label-reasoning consistency check。使用 NLI (Natural Language Inference) 模型 '
        '或 keyword-based heuristics 檢查 reasoning 文字是否支持所給 label。'
        '報告 reasoning_faithfulness_rate。'
    )

    doc.add_heading('4.5 Cross-Model Comparability', level=2)
    p = doc.add_paragraph()
    p.add_run(
        '不同 LLM (GPT-4, Gemma, Claude) 的 CACR 絕對值不可直接比較，'
        '因為 baseline coherence rate 因模型而異。'
    )
    p = doc.add_paragraph()
    p.add_run('建議：').bold = True
    p.add_run(
        '報告 CACR-normalized = (CACR_model - CACR_random) / (1 - CACR_random)。'
        '同時要求每個實驗記錄 model card fields (model name, temperature, max_tokens) '
        '及 variance metrics (CACR std across seeds)。'
    )

    # ============================================================
    # 5. 時空驗證缺口
    # ============================================================
    doc.add_heading('5. 時空驗證缺口 (P1-P2, 水資源 + 社科專家)', level=1)

    doc.add_heading('5.1 空間驗證', level=2)
    p = doc.add_paragraph()
    p.add_run(
        '目前 C&V Module 缺少空間維度驗證。以下為建議新增的 spatial metrics：'
    )
    spatial = [
        ('Adaptation Moran\'s I', '空間自相關指標，檢測適應行為是否具空間聚集性。'
         '以 flood depth 為 spatial weight，計算 adaptation binary (yes/no) 的 Moran\'s I。'
         '預期值 > 0.1 (顯著正自相關)。'),
        ('Flood Zone Gradient Slope', '適應率應隨 flood zone risk 等級遞增。'
         '計算 adaptation_rate vs. flood_zone_risk 的線性斜率，預期 > 0。'),
        ('Buyout Spatial Concentration', '以 Gini coefficient by flood depth 衡量 buyout 的空間集中度。'
         '高風險區應有更高 buyout 率。預期 Gini > 0.3。'),
    ]
    for name, desc in spatial:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f'{name}：')
        run.bold = True
        p.add_run(desc)

    doc.add_heading('5.2 時間驗證', level=2)
    p = doc.add_paragraph()
    p.add_run('建議新增的 temporal metrics：')

    temporal = [
        ('Post-flood Adaptation Spike Ratio',
         '洪災後第一年適應率 / 前三年平均。'
         '預期 1.5-3.0x (Gallagher 2014)。'),
        ('Insurance Survival Half-life',
         '自購買保險後，維持投保的中位年數。'
         '預期 2-4 年 (Michel-Kerjan 2012)。'),
        ('Adaptation S-curve R²',
         '以 logistic function 擬合累積適應率。'
         'R² > 0.85 表示合理的 adoption diffusion pattern。'),
        ('CACR-by-year Slope',
         '偵測 LLM construct coherence 隨時間退化。'
         'Slope < -0.02/year 為 warning signal。'),
    ]
    for name, desc in temporal:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f'{name}：')
        run.bold = True
        p.add_run(desc)

    # ============================================================
    # 6. 社會動態驗證
    # ============================================================
    doc.add_heading('6. 社會動態驗證 (P2, 社科專家)', level=1)
    p = doc.add_paragraph()
    p.add_run(
        '社科專家特別指出，LLM-ABM 的社會動態驗證不應僅依賴個體層級指標，'
        '還需要群體層級的 emergent pattern 檢測：'
    )

    social = [
        ('Spatial Autocorrelation (Moran\'s I)',
         '同一 flood zone 內 agents 的行為是否顯示空間聚集。'
         '若 Moran\'s I 顯著但機制僅為 shared flood exposure (非 social influence)，'
         '需謹慎解讀 — 這是 socially-mediated adaptive behavior，不是 emergent social dynamics。'),
        ('Temporal Contagion (Neighbor Adoption Effect)',
         '鄰居在 t-1 年採取適應行動是否增加 agent 在 t 年的適應機率。'
         '使用 logistic regression 控制 flood exposure，報告 neighbor_effect coefficient 及 p-value。'),
        ('Norm Emergence (Within-group Entropy over Time)',
         '若社會規範正在形成，同一群體內的行為 entropy 應隨時間下降。'
         '注意：WAGF 目前資料顯示 entropy 無明顯下降趨勢，'
         '故不宜 claim "norm emergence"。'),
        ('Rejection Cascade Metric',
         '治理系統 (MG) rejection 導致的連鎖 inaction 效應。'
         '計算 consecutive rejection 後 agent 連續選擇 do_nothing 的平均次數。'
         '若 > 3 次，表示 rejection cascade，需要 governance rule 調整。'),
    ]
    for name, desc in social:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f'{name}：')
        run.bold = True
        p.add_run(desc)

    # ============================================================
    # 7. 可擴充性 (10K+ agents)
    # ============================================================
    doc.add_heading('7. 可擴充性 (10K+ agents, CS 教授)', level=1)
    p = doc.add_paragraph()
    p.add_run(
        '目前模組將所有 trace data 載入記憶體，對 10K+ agents 的大規模實驗不可行。'
        'CS 教授提出以下架構改進：'
    )

    scalability = [
        ('Streaming TraceReader',
         '逐 agent 讀取 trace，記憶體消耗為 O(N_agents) 而非 O(N_decisions)。'
         '使用 generator pattern，避免一次載入所有 JSON。'),
        ('Single-pass Benchmark Extraction',
         '所有 benchmark metrics 在一次 trace scan 中同時計算，'
         '避免重複讀取同一份 trace 資料。'),
        ('Per-seed Parallelism',
         '使用 ProcessPoolExecutor 平行處理不同 seed 的驗證。'
         '每個 seed 獨立計算，最後合併統計。'),
        ('orjson for JSON Parsing',
         '替換標準 json module 為 orjson，可獲得 3-5x parsing 加速。'
         '特別對大量小 JSON 文件效果顯著。'),
    ]
    for name, desc in scalability:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f'{name}：')
        run.bold = True
        p.add_run(desc)

    p = doc.add_paragraph()
    p.add_run('預估效果：').bold = True
    p.add_run(
        '500K traces 情境下，從 ~4GB 記憶體消耗降至 ~50MB，'
        '處理時間從 O(N²) 降至 O(N)。'
    )

    # ============================================================
    # 8. 跨領域擴展
    # ============================================================
    doc.add_heading('8. 跨領域擴展 (水資源專家)', level=1)
    p = doc.add_paragraph()
    p.add_run(
        'C&V Module 目前針對 flood adaptation 設計，但 WAGF 的願景是成為通用 LLM-ABM 框架。'
        '水資源專家提出以下跨領域擴展方向：'
    )

    domains = [
        ('Groundwater Management',
         'Decision-Analytic (DA) / Prospect-Framing (PF) constructs。'
         '治理規則對應 SGMA (Sustainable Groundwater Management Act) 法規。'
         'Benchmarks: pumping reduction rate, basin sustainability index。'),
        ('Urban Water Demand',
         'Conservation motivation constructs (Environmental concern, Social norm, Economic incentive)。'
         'Benchmarks: per-capita demand reduction, price elasticity of demand (-0.1 to -0.6)。'),
        ('Drought Response',
         'Stress-coping constructs for agricultural agents。'
         'Benchmarks: fallowing rates, water market trading volumes, crop switching frequency。'),
        ('Irrigation (已有實作)',
         'WSA/ACA dual-appraisal constructs 已在 irrigation ABM 中實作。'
         '但缺少正式的 ConstructCoherenceSpec — 目前 coherence 定義散佈在程式碼中，'
         '需要整理為 YAML 格式的 spec。'),
    ]
    for name, desc in domains:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f'{name}：')
        run.bold = True
        p.add_run(desc)

    # ============================================================
    # 9. Elevation Rate = 0.57 問題
    # ============================================================
    doc.add_heading('9. Elevation Rate = 0.57 問題 (水資源專家)', level=1)
    p = doc.add_paragraph()
    p.add_run(
        '目前 flood ABM 的 elevation rate 為 0.57，遠超過 empirical benchmark 範圍 (0.10-0.35)。'
        '水資源專家分析其根本原因及建議：'
    )

    doc.add_heading('9.1 根因分析', level=2)
    causes = [
        ('Benchmark 範圍合理但窄',
         '0.10-0.35 的範圍基於全國平均數據，實際因地區而異。'
         '但 0.57 即使在高風險區也顯著偏高。'),
        ('缺少 Capacity Constraint',
         '模型未考慮 contractor throughput limit — 現實中每年只有有限的承包商可執行 elevation 工程。'
         '即使所有 owner 都想 elevate，供給面的限制也會壓低實際 elevation rate。'),
        ('Subsidy Feedback Loop 無 Damping',
         '當 elevation rate 高時，更多 agents 看到鄰居 elevate → 更多人 elevate → positive feedback。'
         '現實中，subsidy 預算有限，會自然 dampen 這個 loop。'),
    ]
    for name, desc in causes:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f'{name}：')
        run.bold = True
        p.add_run(desc)

    doc.add_heading('9.2 建議修正', level=2)
    fixes = [
        '加入 per-year elevation cap (5-8% of unelevated owners)，模擬 contractor capacity constraint。',
        '報告 subsidy-conditioned elevation rate (有補助 vs. 無補助)，分離 economic motivation vs. risk motivation。',
        '在 benchmark comparison 中加入 capacity-adjusted expected rate。',
    ]
    for fix in fixes:
        doc.add_paragraph(fix, style='List Bullet')

    # ============================================================
    # 10. 實施路線圖
    # ============================================================
    doc.add_heading('10. 實施路線圖', level=1)
    p = doc.add_paragraph()
    p.add_run(
        '以下為建議的分階段實施計畫，總計約 9 個工作天：'
    )

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    header_cells = table.rows[0].cells
    headers = ['Phase', '內容', '工時', '依賴']
    for i, h in enumerate(headers):
        p = header_cells[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(255, 255, 255)
        tc = header_cells[i]._element
        tcPr = tc.get_or_add_tcPr()
        shading = tcPr.makeelement(qn('w:shd'), {
            qn('w:fill'): '4472C4',
            qn('w:val'): 'clear',
        })
        tcPr.append(shading)

    phases = [
        ('Phase 0', '修 P0 bugs (EBE, TP/CP sentinel, agent type join)\n+ golden regression test', '0.5 天', '無'),
        ('Phase 1', '常數外部化 YAML\n(rules, benchmarks, coherence tables)', '1 天', 'Phase 0'),
        ('Phase 2', '拆分子模組\n(metrics/, io/, reporting/)', '2 天', 'Phase 1'),
        ('Phase 3', 'BehavioralTheory protocol\n+ TheoryRegistry', '2 天', 'Phase 2'),
        ('Phase 4', 'BenchmarkComputation plugins\n+ YAML 定義', '2 天', 'Phase 2'),
        ('Phase 5', 'Streaming TraceReader\n+ ValidationRunner facade', '1.5 天', 'Phase 2-4'),
        ('Total', '', '~9 天', ''),
    ]
    for row_data in phases:
        row = table.add_row()
        for i, val in enumerate(row_data):
            p = row.cells[i].paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(9)
            if row_data[0] == 'Total':
                run.bold = True

    # Set column widths
    for row in table.rows:
        row.cells[0].width = Cm(2.0)
        row.cells[1].width = Cm(8.0)
        row.cells[2].width = Cm(2.0)
        row.cells[3].width = Cm(3.0)

    doc.add_paragraph()  # spacer

    p = doc.add_paragraph()
    p.add_run('Phase 優先級說明：').bold = True
    p.add_run(
        'Phase 0 為 blocking — 必須在任何 production run 之前完成。'
        'Phase 1-2 為基礎架構，Phase 3-5 可平行進行。'
        '建議在 Phase 0 完成後即可繼續 400×13yr production run，'
        '架構重構可與 production run 平行進行。'
    )

    # ============================================================
    # 11. 附錄
    # ============================================================
    doc.add_heading('11. 附錄：各專家報告摘要', level=1)

    doc.add_heading('11.1 LLM 行為工程專家', level=2)
    p = doc.add_paragraph()
    p.add_run('15 項建議，依 priority 分佈：')

    llm_table = doc.add_table(rows=1, cols=3)
    llm_table.style = 'Table Grid'
    headers = ['Priority', '數量', '主要項目']
    for i, h in enumerate(headers):
        p = llm_table.rows[0].cells[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(255, 255, 255)
        tc = llm_table.rows[0].cells[i]._element
        tcPr = tc.get_or_add_tcPr()
        shading = tcPr.makeelement(qn('w:shd'), {
            qn('w:fill'): '4472C4',
            qn('w:val'): 'clear',
        })
        tcPr.append(shading)

    llm_rows = [
        ('P0', '2', 'EBE entropy bug, TP/CP silent default'),
        ('P1', '3', 'Sycophancy detection, CACR temporal, Construct grounding'),
        ('P2', '6', 'Reasoning faithfulness, Cross-model normalization,\nPrompt-response independence, Label circularity,\nHallucination taxonomy, Output format drift'),
        ('P3', '4', 'Confidence calibration, Chain-of-thought analysis,\nMulti-turn consistency, Temperature sensitivity'),
    ]
    for row_data in llm_rows:
        row = llm_table.add_row()
        for i, val in enumerate(row_data):
            run = row.cells[i].paragraphs[0].add_run(val)
            run.font.size = Pt(9)

    doc.add_heading('11.2 行為/社會科學專家', level=2)
    p = doc.add_paragraph()
    p.add_run('8 項建議，依 severity 排列：')
    social_items = [
        '(Critical) Construct label circularity — self-consistency ≠ construct validity',
        '(Critical) Agent type inference circularity',
        '(High) Frame-conditional theory support (PT, Nudge)',
        '(High) Social contagion metric — neighbor adoption effect',
        '(Medium) Norm emergence entropy — 需配合 "no convergence claim" 原則',
        '(Medium) Rejection cascade metric',
        '(Low) Cross-cultural construct validity (not applicable to current study)',
        '(Low) Longitudinal construct stability check',
    ]
    for item in social_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('11.3 水資源工程專家', level=2)
    p = doc.add_paragraph()
    p.add_run('8 項建議，依 priority 排列：')
    water_items = [
        '(P0) Elevation rate capacity constraint — per-year cap',
        '(P1) Spatial validation metrics (Moran\'s I, gradient slope)',
        '(P1) Post-flood spike ratio benchmark',
        '(P1) Insurance survival half-life',
        '(P2) Cross-domain extension specs (groundwater, urban, drought)',
        '(P2) Subsidy feedback damping',
        '(P3) Climate scenario sensitivity (RCP/SSP pathways)',
        '(P3) Compound hazard interaction (flood + drought)',
    ]
    for item in water_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('11.4 資工系教授', level=2)
    p = doc.add_paragraph()
    p.add_run('完整架構重構方案：')
    cs_items = [
        '7 子模組拆分 (theories, hallucinations, benchmarks, metrics, io, reporting, engine)',
        '5 phase migration plan (見 Section 10)',
        'Streaming TraceReader with O(N_agents) memory',
        'Plugin-based benchmark computation',
        '60+ test strategy: unit tests per sub-module + integration test suite + regression golden files',
        'CI/CD integration: validation metrics as GitHub Actions check',
    ]
    for item in cs_items:
        doc.add_paragraph(item, style='List Bullet')

    # ============================================================
    # Footer
    # ============================================================
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('— End of Report —')
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    run.italic = True

    # Save
    doc.save(OUTPUT)
    print(f"Document saved to: {OUTPUT}")


if __name__ == '__main__':
    build_document()
