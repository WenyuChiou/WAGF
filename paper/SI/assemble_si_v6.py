"""
Assemble SAGE_WRR_SI_Tables_v6.docx from individual markdown section files.

Structure:
  Part A: Framework-Level (S1, S2)
  Part B: Flood Case Study (S3=old Table S1/S2, S4=old S7 diagnostics, S5=figures)
  Part C: Irrigation Case Study (S6=governance rules, S7=FQL mapping, S8=mass balance, S9=old Table S3)
"""
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

SI_DIR = Path(__file__).parent
PAPER_DIR = SI_DIR.parent
OUT_PATH = PAPER_DIR / "SAGE_WRR_SI_Tables_v6.docx"

# ── helpers ──────────────────────────────────────────────────────────────────

def set_cell_text(cell, text, bold=False, size=10):
    """Set cell text with formatting."""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold


def add_table_from_md(doc, lines, first_row_bold=True):
    """Parse markdown table lines and add a Word table."""
    # Filter to only lines that start with |
    table_lines = [l for l in lines if l.strip().startswith("|")]
    if len(table_lines) < 2:
        return

    # Parse header
    header = [c.strip() for c in table_lines[0].strip().strip("|").split("|")]

    # Find data rows (skip separator row)
    data_rows = []
    for line in table_lines[1:]:
        stripped = line.strip().strip("|")
        # Skip separator rows (all dashes/colons/spaces)
        if all(c in "-: |" for c in stripped):
            continue
        data_rows.append([c.strip() for c in stripped.split("|")])

    ncols = len(header)
    nrows = 1 + len(data_rows)

    table = doc.add_table(rows=nrows, cols=ncols)
    table.style = "Table Grid"
    table.autofit = True

    # Header row
    for j, h in enumerate(header):
        if j < ncols:
            set_cell_text(table.rows[0].cells[j], h, bold=True, size=10)

    # Data rows
    for i, row_data in enumerate(data_rows):
        for j, val in enumerate(row_data):
            if j < ncols:
                # Clean markdown formatting
                val = val.replace("`", "").replace("**", "").replace("*", "")
                val = re.sub(r'\$[^$]+\$', lambda m: m.group(0).strip("$"), val)
                set_cell_text(table.rows[i + 1].cells[j], val, bold=False, size=10)

    doc.add_paragraph()  # spacing after table


def add_heading(doc, text, level):
    """Add a heading with Times New Roman font."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(0, 0, 0)


def add_body(doc, text):
    """Add a body paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 2.0
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    return p


def add_code_block(doc, lines):
    """Add a code block as formatted paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(1.0)
    text = "\n".join(lines)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def process_md_file(doc, filepath, renumber_prefix=None):
    """
    Read a markdown file and add its content to the Word document.
    Handles headings, paragraphs, tables, and code blocks.
    """
    text = filepath.read_text(encoding="utf-8")
    lines = text.split("\n")

    i = 0
    in_code_block = False
    code_lines = []
    table_buffer = []

    while i < len(lines):
        line = lines[i]

        # Code block toggle
        if line.strip().startswith("```"):
            if in_code_block:
                # End code block
                add_code_block(doc, code_lines)
                code_lines = []
                in_code_block = False
            else:
                # Flush any table buffer first
                if table_buffer:
                    add_table_from_md(doc, table_buffer)
                    table_buffer = []
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Table lines
        if line.strip().startswith("|"):
            table_buffer.append(line)
            i += 1
            continue
        else:
            if table_buffer:
                add_table_from_md(doc, table_buffer)
                table_buffer = []

        # Skip empty lines
        if not line.strip():
            i += 1
            continue

        # LaTeX display math ($$...$$)
        if line.strip().startswith("$$"):
            # Collect until closing $$
            math_lines = [line.strip().lstrip("$")]
            if not line.strip().endswith("$$") or line.strip() == "$$":
                i += 1
                while i < len(lines) and not lines[i].strip().endswith("$$"):
                    math_lines.append(lines[i].strip())
                    i += 1
                if i < len(lines):
                    closing = lines[i].strip().rstrip("$")
                    if closing:
                        math_lines.append(closing)
            else:
                math_lines = [line.strip().strip("$")]

            math_text = " ".join(l for l in math_lines if l)
            # Simplify LaTeX for Word display
            math_text = math_text.replace("\\text{", "").replace("\\bigl(", "(").replace("\\bigr)", ")")
            math_text = math_text.replace("\\;", " ").replace("\\,", "").replace("\\!", "")
            math_text = math_text.replace("\\quad", "    ").replace("\\qquad", "        ")
            math_text = math_text.replace("\\frac{", "(").replace("}{", ")/(")
            math_text = math_text.replace("\\left(", "(").replace("\\right)", ")")
            math_text = math_text.replace("\\min", "min").replace("\\max", "max")
            math_text = math_text.replace("\\sum_", "SUM_").replace("\\in", " in ")
            math_text = math_text.replace("\\geq", ">=").replace("\\leq", "<=")
            math_text = math_text.replace("\\times", "x").replace("\\cdot", ".")
            math_text = re.sub(r'\\mathcal\{([^}]+)\}', r'\1', math_text)
            math_text = re.sub(r'\\operatorname\{([^}]+)\}', r'\1', math_text)
            # Clean remaining braces
            math_text = math_text.replace("{", "").replace("}", "")

            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(math_text)
            run.font.name = "Cambria Math"
            run.font.size = Pt(11)
            run.italic = True
            i += 1
            continue

        # Headings
        if line.startswith("#"):
            hashes = len(line) - len(line.lstrip("#"))
            title = line.lstrip("#").strip()
            # Map heading levels
            level = min(hashes, 4)
            add_heading(doc, title, level=level)
            i += 1
            continue

        # Bold/italic lines
        if line.startswith("**") and line.endswith("**"):
            p = doc.add_paragraph()
            run = p.add_run(line.strip("*").strip())
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
            run.bold = True
            i += 1
            continue

        # Numbered list items
        if re.match(r'^\d+\.', line.strip()):
            text_content = line.strip()
            # Clean markdown
            text_content = text_content.replace("`", "").replace("**", "")
            text_content = re.sub(r'\$[^$]+\$', lambda m: m.group(0).strip("$"), text_content)
            p = doc.add_paragraph(style="List Number")
            run = p.add_run(text_content)
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
            i += 1
            continue

        # Bullet list items
        if line.strip().startswith("- ") or line.strip().startswith("* "):
            text_content = line.strip().lstrip("-* ").strip()
            text_content = text_content.replace("`", "").replace("**", "")
            text_content = re.sub(r'\$[^$]+\$', lambda m: m.group(0).strip("$"), text_content)
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(text_content)
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
            i += 1
            continue

        # Regular paragraph
        text_content = line.strip()
        text_content = text_content.replace("`", "")
        # Clean inline bold
        text_content = re.sub(r'\*\*([^*]+)\*\*', r'\1', text_content)
        # Clean inline math
        text_content = re.sub(r'\$([^$]+)\$', r'\1', text_content)
        if text_content:
            add_body(doc, text_content)
        i += 1

    # Flush remaining buffers
    if table_buffer:
        add_table_from_md(doc, table_buffer)
    if code_lines:
        add_code_block(doc, code_lines)


def add_figure_placeholder(doc, fig_id, caption):
    """Add a figure placeholder with caption."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"[{fig_id}: See separate image file]")
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    run.italic = True

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(caption)
    run2.font.name = "Times New Roman"
    run2.font.size = Pt(10)
    run2.bold = True
    doc.add_paragraph()


# ── main assembly ────────────────────────────────────────────────────────────

def main():
    doc = Document()

    # ── Page setup ──
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # ── Default font ──
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    style.paragraph_format.line_spacing = 2.0

    # ── Title ──
    title = doc.add_heading("Supporting Information for", level=1)
    for run in title.runs:
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(0, 0, 0)

    subtitle = doc.add_heading(
        "WAGF: A Governance Framework for LLM-Driven Agent-Based Modeling "
        "in Human-Water Systems", level=2
    )
    for run in subtitle.runs:
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(0, 0, 0)

    doc.add_paragraph()

    # ── Table of Contents ──
    add_heading(doc, "Contents", level=2)
    toc_items = [
        "Part A: Framework-Level",
        "  S1. Prompt Templates and Response Format",
        "  S2. Governance Retry Mechanism and EarlyExit",
        "Part B: Flood Case Study",
        "  S3. Flood Cross-Model Metrics (Table S1, Table S2)",
        "  S4. Flood Behavioral Diagnostics",
        "  S5. Flood Adaptation Matrices (Figures S1-S5)",
        "Part C: Irrigation Case Study",
        "  S6. Complete Governance Rule Specification",
        "  S7. FQL-to-LLM Persona Cluster Mapping",
        "  S8. Mass Balance and Human-Water Coupling",
        "  S9. Irrigation Governance Summary (Table S6)",
    ]
    for item in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(item)
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
        if not item.startswith("  "):
            run.bold = True

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════
    # PART A: Framework-Level
    # ════════════════════════════════════════════════════════════════════
    add_heading(doc, "Part A: Framework-Level", level=1)

    # S1
    process_md_file(doc, SI_DIR / "Section_S1_Prompt_Templates.md")
    doc.add_page_break()

    # S2
    process_md_file(doc, SI_DIR / "Section_S2_Retry_EarlyExit.md")
    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════
    # PART B: Flood Case Study
    # ════════════════════════════════════════════════════════════════════
    add_heading(doc, "Part B: Flood Case Study", level=1)

    # S3 — placeholder for existing Table S1/S2 (from v5.docx, to be pasted manually)
    add_heading(doc, "S3. Flood Cross-Model Metrics", level=2)
    add_body(doc,
        "Table S1: Complete 18-row multi-model data (6 models x 3 groups): "
        "R_R, H_norm, EHE, FF. [Retained from SAGE_WRR_SI_Tables_v5.docx]")
    add_body(doc,
        "Table S2: R_H (strict feasibility safety diagnostic) by model size. "
        "[Retained from SAGE_WRR_SI_Tables_v5.docx]")
    doc.add_paragraph()

    # S4 — Behavioral Diagnostics (old S7)
    process_md_file(doc, SI_DIR / "Section_S7_Behavioral_Diagnostics_Examples.md")
    doc.add_page_break()

    # S5 — Flood figures (placeholders)
    add_heading(doc, "S5. Flood Adaptation Matrices", level=2)
    fig_placeholders = [
        ("Figure S1", "6x3 Adaptation matrix across all models and governance groups"),
        ("Figure S2", "Cumulative relocation curves (A=0%, B=32%, C=37%)"),
        ("Figure S3", "Economic hallucination fix comparison (v4 vs v6)"),
        ("Figure S4", "Gemma3 3x3 adaptation matrix"),
        ("Figure S5", "Ministral3 3x3 adaptation matrix"),
    ]
    for fig_id, caption in fig_placeholders:
        add_figure_placeholder(doc, fig_id, f"{fig_id}. {caption}")

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════
    # PART C: Irrigation Case Study
    # ════════════════════════════════════════════════════════════════════
    add_heading(doc, "Part C: Irrigation Case Study", level=1)

    # S6 — Governance Rules (new)
    process_md_file(doc, SI_DIR / "Section_S6_Governance_Rules.md")
    doc.add_page_break()

    # S7 — FQL Cluster Mapping (new)
    process_md_file(doc, SI_DIR / "Section_S7_FQL_Cluster_Mapping.md")
    doc.add_page_break()

    # S8 — Mass Balance (updated)
    process_md_file(doc, SI_DIR / "Section_S8_Mass_Balance.md")
    doc.add_page_break()

    # S9 — Irrigation Governance Summary (old Table S3)
    add_heading(doc, "S9. Irrigation Governance Summary", level=2)
    process_md_file(doc, SI_DIR / "Table_S3_Irrigation_Governance.md")

    # ── Save ──
    doc.save(str(OUT_PATH))
    print(f"Created: {OUT_PATH}")
    print(f"Sections: S1-S9 (4 new + 5 existing/renumbered)")


if __name__ == "__main__":
    main()
