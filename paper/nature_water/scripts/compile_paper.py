#!/usr/bin/env python3
"""
Compile all Nature Water paper sections into a single formatted Word document.

Nature Water Analysis format:
- Title + authors
- Abstract (≤150 words, no references)
- Main text (~4,000 words): Introduction, Results, Discussion
- Methods
- References
- Supplementary Information (separate file)

Formatting: Times New Roman 11pt, justified, double-spaced, line numbers.
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# Paths
DRAFTS = Path(r"C:\Users\wenyu\Desktop\Lehigh\governed_broker_framework\paper\nature_water\drafts")
OUTPUT = Path(r"C:\Users\wenyu\Desktop\Lehigh\governed_broker_framework\paper\nature_water")

FONT_NAME = "Times New Roman"
FONT_SIZE = Pt(11)
HEADING1_SIZE = Pt(14)
HEADING2_SIZE = Pt(12)
TABLE_FONT_SIZE = Pt(9)
LINE_SPACING = WD_LINE_SPACING.DOUBLE


FIRST_LINE_INDENT = Inches(0.3)  # First-line indent for body paragraphs
FIGURES_DIR = Path(r"C:\Users\wenyu\Desktop\Lehigh\governed_broker_framework\paper\nature_water\figures")


def set_paragraph_format(para, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                         space_after=Pt(0), space_before=Pt(0),
                         line_spacing=LINE_SPACING,
                         first_line_indent=None):
    """Apply standard formatting to a paragraph."""
    pf = para.paragraph_format
    pf.alignment = alignment
    pf.space_after = space_after
    pf.space_before = space_before
    pf.line_spacing_rule = line_spacing
    if first_line_indent is not None:
        pf.first_line_indent = first_line_indent


def set_run_font(run, size=FONT_SIZE, bold=False, italic=False, color=None):
    """Apply font settings to a run."""
    run.font.name = FONT_NAME
    run.font.size = size
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    # Ensure CJK font compatibility
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{FONT_NAME}"/>')
        rPr.append(rFonts)
    else:
        rFonts.set(qn('w:eastAsia'), FONT_NAME)


def add_formatted_text(para, text, size=FONT_SIZE, bold=False, italic=False):
    """Add text with formatting, handling **bold** and *italic* markers."""
    # Split on bold markers first
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            inner = part[2:-2]
            # Check for italic within bold
            italic_parts = re.split(r'(\*[^*]+\*)', inner)
            for ip in italic_parts:
                if ip.startswith('*') and ip.endswith('*'):
                    run = para.add_run(ip[1:-1])
                    set_run_font(run, size, bold=True, italic=True)
                else:
                    run = para.add_run(ip)
                    set_run_font(run, size, bold=True, italic=False)
        else:
            # Check for italic
            italic_parts = re.split(r'(\*[^*]+\*)', part)
            for ip in italic_parts:
                if ip.startswith('*') and ip.endswith('*'):
                    run = para.add_run(ip[1:-1])
                    set_run_font(run, size, bold=bold, italic=True)
                else:
                    run = para.add_run(ip)
                    set_run_font(run, size, bold=bold, italic=italic)


def add_heading(doc, text, level=1):
    """Add a heading with NW formatting."""
    para = doc.add_paragraph()
    size = HEADING1_SIZE if level == 1 else HEADING2_SIZE
    run = para.add_run(text)
    set_run_font(run, size=size, bold=True)
    set_paragraph_format(para, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                         space_before=Pt(18), space_after=Pt(6))
    return para


def add_body_paragraph(doc, text, indent=True):
    """Add a body paragraph with justified alignment, formatted text, and first-line indent."""
    para = doc.add_paragraph()
    add_formatted_text(para, text)
    set_paragraph_format(para, first_line_indent=FIRST_LINE_INDENT if indent else None)
    return para


def parse_markdown_table(lines):
    """Parse markdown table lines into header + rows."""
    table_lines = [l.strip() for l in lines if l.strip() and not re.match(r'^\|[-:| ]+\|$', l.strip())]
    if not table_lines:
        return [], []
    header = [c.strip() for c in table_lines[0].split('|')[1:-1]]
    rows = []
    for line in table_lines[1:]:
        cells = [c.strip() for c in line.split('|')[1:-1]]
        rows.append(cells)
    return header, rows


def add_table(doc, header, rows, caption=None):
    """Add a formatted table to the document."""
    if caption:
        para = doc.add_paragraph()
        add_formatted_text(para, caption, size=Pt(10), bold=True)
        set_paragraph_format(para, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                             space_before=Pt(12), space_after=Pt(4))

    n_cols = len(header)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    for i, h in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.text = ""
        para = cell.paragraphs[0]
        run = para.add_run(h)
        set_run_font(run, size=TABLE_FONT_SIZE, bold=True)
        set_paragraph_format(para, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                             line_spacing=WD_LINE_SPACING.SINGLE)
        # Gray background for header
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="E8E8E8" w:val="clear"/>')
        cell._element.get_or_add_tcPr().append(shading)

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            if c_idx >= n_cols:
                break
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            para = cell.paragraphs[0]
            # Handle bold in cells
            add_formatted_text(para, cell_text, size=TABLE_FONT_SIZE)
            align = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
            set_paragraph_format(para, alignment=align,
                                 line_spacing=WD_LINE_SPACING.SINGLE)

    return table


def extract_body_and_tables(md_text):
    """
    Parse markdown into a sequence of (type, content) tuples.
    Types: 'heading1', 'heading2', 'heading3', 'paragraph', 'table', 'table_note'
    """
    elements = []
    lines = md_text.split('\n')
    i = 0
    # Skip front matter (lines starting with # that are metadata)
    while i < len(lines):
        line = lines[i].strip()
        if line.startswith('# ') and not line.startswith('# Nature Water'):
            # Check if it's a section heading vs metadata
            if line in ('# Introduction', '# Results', '# Discussion', '# Methods',
                        '# Supplementary Information') or line.startswith('## ') or line.startswith('### '):
                break
        if line == '---':
            i += 1
            # Skip until we find actual content
            while i < len(lines) and not lines[i].strip():
                i += 1
            continue
        if line.startswith('## Date:') or line.startswith('## ≤') or line.startswith('## Structure') or line.startswith('## Narrative') or line.startswith('## Status') or line.startswith('## Data:'):
            i += 1
            continue
        if line.startswith('# Nature Water') or line.startswith('# Supplementary'):
            i += 1
            continue
        if not line:
            i += 1
            continue
        break

    # Now parse content
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Skip metadata / changelog
        if stripped.startswith('## Structural changes') or stripped.startswith('## Changes from') or stripped.startswith('### Word count') or stripped.startswith('### Changes') or stripped.startswith('### R2 polish') or stripped.startswith('### Review history') or stripped.startswith('### v10'):
            # Skip until end of file or next ---
            while i < len(lines) and lines[i].strip() != '---':
                i += 1
            i += 1
            continue

        if stripped == '---':
            i += 1
            continue

        # Headings
        if stripped.startswith('### '):
            text = stripped[4:].strip()
            if text and not text.startswith('Word count') and not text.startswith('Changes'):
                elements.append(('heading3', text))
            i += 1
            continue
        if stripped.startswith('## '):
            text = stripped[3:].strip()
            if text and not text.startswith('Date:') and not text.startswith('≤') and not text.startswith('Structural'):
                elements.append(('heading2', text))
            i += 1
            continue
        if stripped.startswith('# '):
            text = stripped[2:].strip()
            if text in ('Introduction', 'Results', 'Discussion', 'Methods'):
                elements.append(('heading1', text))
            i += 1
            continue

        # Tables
        if stripped.startswith('|') and i + 1 < len(lines) and '---' in lines[i + 1]:
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            header, rows = parse_markdown_table(table_lines)
            if header:
                elements.append(('table', (header, rows)))
            continue

        # Table caption (bold paragraph starting with Table)
        if stripped.startswith('**Table '):
            elements.append(('table_caption', stripped))
            i += 1
            continue

        # Table notes (italic, starting with *)
        if stripped.startswith('*') and not stripped.startswith('**'):
            elements.append(('table_note', stripped))
            i += 1
            continue

        # Numbered list (Methods pipeline)
        if re.match(r'^\d+\.', stripped):
            elements.append(('list_item', stripped))
            i += 1
            continue

        # Bullet list
        if stripped.startswith('- '):
            elements.append(('bullet', stripped[2:]))
            i += 1
            continue

        # Regular paragraph
        if stripped:
            # Collect multi-line paragraph
            para_lines = [stripped]
            i += 1
            while i < len(lines):
                next_line = lines[i].strip()
                if not next_line or next_line.startswith('#') or next_line.startswith('|') or next_line.startswith('**Table') or next_line.startswith('*') or next_line.startswith('- ') or next_line == '---' or re.match(r'^\d+\.', next_line):
                    break
                para_lines.append(next_line)
                i += 1
            full_para = ' '.join(para_lines)
            elements.append(('paragraph', full_para))
            continue

        i += 1

    return elements


def add_figure(doc, image_path, caption, width=Inches(6.5)):
    """Add a figure with caption to the document."""
    # Image
    para = doc.add_paragraph()
    set_paragraph_format(para, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                         space_before=Pt(12), space_after=Pt(4))
    run = para.add_run()
    run.add_picture(str(image_path), width=width)

    # Caption
    cap_para = doc.add_paragraph()
    add_formatted_text(cap_para, caption, size=Pt(9))
    set_paragraph_format(cap_para, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                         space_before=Pt(2), space_after=Pt(12))
    return cap_para


def add_line_numbers(doc):
    """Add line numbers to the document (NW submission requirement)."""
    for section in doc.sections:
        sectPr = section._sectPr
        ln_num = parse_xml(f'<w:lnNumType {nsdecls("w")} w:countBy="1" w:restart="continuous"/>')
        sectPr.append(ln_num)


def compile_main_paper():
    """Compile the main paper (Abstract + Intro + Results + Discussion + Methods)."""
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    # Default font
    style = doc.styles['Normal']
    style.font.name = FONT_NAME
    style.font.size = FONT_SIZE
    style.paragraph_format.line_spacing_rule = LINE_SPACING
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # ==================== TITLE ====================
    para = doc.add_paragraph()
    run = para.add_run("Institutional Constraints Widen Adaptive Behavioural Diversity\nin Language-Based Water Agents")
    set_run_font(run, size=Pt(16), bold=True)
    set_paragraph_format(para, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                         space_before=Pt(48), space_after=Pt(24))

    # Authors
    para = doc.add_paragraph()
    run = para.add_run("Wen-Yuan Chen¹*, Y. C. Ethan Yang¹")
    set_run_font(run, size=Pt(12))
    set_paragraph_format(para, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                         space_after=Pt(6))

    # Affiliation
    para = doc.add_paragraph()
    run = para.add_run("¹ Department of Civil and Environmental Engineering, Lehigh University, Bethlehem, PA, USA")
    set_run_font(run, size=Pt(10), italic=True)
    set_paragraph_format(para, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                         space_after=Pt(4))

    para = doc.add_paragraph()
    run = para.add_run("* Corresponding author: wec225@lehigh.edu")
    set_run_font(run, size=Pt(10), italic=True)
    set_paragraph_format(para, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                         space_after=Pt(24))

    # ==================== ABSTRACT ====================
    add_heading(doc, "Abstract", level=1)

    abstract_md = (DRAFTS / "abstract_v10.md").read_text(encoding='utf-8')
    # Extract just the abstract body (between the --- markers)
    match = re.search(r'---\s*\n(.*?)\n---', abstract_md, re.DOTALL)
    if match:
        abstract_text = match.group(1).strip()
        # Remove markdown formatting markers from front matter
        abstract_lines = []
        for line in abstract_text.split('\n'):
            line = line.strip()
            if line and not line.startswith('#') and not line.startswith('##'):
                abstract_lines.append(line)
        abstract_body = ' '.join(abstract_lines)
        add_body_paragraph(doc, abstract_body)

    # ==================== MAIN SECTIONS ====================
    sections_files = [
        "introduction_v10.md",
        "section2_v11_results.md",
        "section3_v11_discussion.md",
    ]

    # Figure insertion map: after which section heading to insert which figure
    # Figures are inserted AFTER processing the section
    # Main text: 4 figures (framework, irrigation, flood, cross-model)
    # Tables moved to Extended Data (not counted toward 4-item limit)
    figure_after_section = {
        "introduction_v10.md": [
            (FIGURES_DIR / "Fig1_framework.png",
             "**Figure 1. Water Agent Governance Framework architecture.** "
             "At each decision step, an LLM agent receives contextual information "
             "(personal state, social observations, system indicators) and proposes "
             "an action with natural-language reasoning. The governance pipeline "
             "validates the proposal through six sequential checks: schema validation, "
             "action legality, physical feasibility, institutional compliance, "
             "magnitude plausibility, and theory consistency. Failed proposals receive "
             "structured feedback for revision (up to three attempts). "
             "All domain-specific knowledge is expressed through declarative configuration "
             "files (action registries, agent-type specifications, persona prompts), "
             "making the architecture transferable across water domains without code changes.",
             Inches(6.0)),
        ],
        "section2_v11_results.md": [
            (FIGURES_DIR / "Fig2_irrigation_v2.png",
             "**Figure 2. Governance shapes adaptive exploitation in the irrigation domain** "
             "(78 agents \u00d7 42 years, Gemma-3 4B, 3 seeds). "
             "(**a**) Action composition over time for three conditions: governed (with demand-ceiling rule, "
             "IBR = 42%), no ceiling (demand-ceiling removed), and ungoverned (all governance removed, "
             "IBR = 91%). Stacked areas show ensemble-mean action shares. "
             "Black solid line: condition-specific Lake Mead level (right axis); "
             "grey dashed line: fuzzy Q-learning (FQL) baseline (Hung and Yang, 2021). "
             "Red dashed line: Tier 1 shortage threshold (1,075 ft); grey shading: drought periods. "
             "IBR = Irrational Behaviour Rate (fraction of high-scarcity decisions proposing demand increases). "
             "(**b**) Demand ratio (requested volume / historical baseline) versus shortage years "
             "(years with mean Lake Mead below 1,075 ft). Each small dot represents one seed; "
             "large markers show condition means. Arrows trace ablation path: removing only the "
             "demand-ceiling rule shifts toward over-extraction; removing all governance collapses extraction. "
             "Governed agents occupy the adaptive exploitation zone (high extraction, moderate shortage). "
             "Shaded bands show \u00b1 1 s.d. across 3 seeds.",
             Inches(6.0)),
            (FIGURES_DIR / "Fig3_pie_v3.png",
             "**Figure 3. Structured non-compliance at institutional boundaries across two water domains.** "
             "(**a**) Irrigation domain: 5\u00d75 pie matrix showing action distributions across "
             "Water Shortage Appraisal (WSA, rows) and Adaptive Capacity Appraisal (ACA, columns) levels "
             "(78 agents \u00d7 42 years, Gemma-3 4B, 3 seeds). Pie sizes proportional to decision count (n). "
             "Non-compliance concentrates at the high-WSA/high-ACA boundary where 54.8% of governed agents "
             "still proposed demand increases (100% rejected by governance rules). "
             "(**b**) Flood domain: 5\u00d75 pie matrix showing cumulative protection state across "
             "Threat Appraisal (TA, rows) and Coping Appraisal (CA, columns) levels "
             "(100 agents \u00d7 10 years, Gemma-3 4B, 3 seeds). "
             "92% of governed decisions self-assessed coping appraisal at a single level (Med CA, highlighted), "
             "suggesting governance success in flood partially reflects compression of cognitive differentiation. "
             "Colours: Okabe-Ito palette consistent across all figures.",
             Inches(6.0)),
            (FIGURES_DIR / "Fig4_crossmodel.png",
             "**Figure 4. Governance effect on behavioural diversity and rationality across six LLMs** "
             "(flood domain, 100 agents \u00d7 10 years, 3 runs per condition). "
             "Models sorted by parameter count (3.2B to 27B). "
             "(**a**) Paired dot plot showing ungoverned (vermillion) versus governed (blue) "
             "behavioural diversity (normalised Shannon entropy H/log\u2082(k), k = 4 action types); "
             "connecting lines indicate effect direction. "
             "(**b**) Forest plot of governance effect (\u0394 = governed \u2212 ungoverned) with 95% CIs; "
             "blue = statistically significant, grey = non-significant. "
             "Diamond shows mean effect across all six models (+0.123). "
             "(**c**) IBR (Irrational Behaviour Rate, %; fraction of physically impossible or "
             "PMT-inconsistent decisions) for each model. Governance reduced IBR in all six models, "
             "with the largest reductions for Ministral 14B (11.6% \u2192 0.4%) and Ministral 3B "
             "(8.9% \u2192 1.7%). Error bars: \u00b1 1 s.d. across 3 seeds. "
             "Full IBR decomposition in Supplementary Table 1.",
             Inches(6.0)),
        ],
    }

    def render_elements(elements):
        """Render parsed markdown elements into the Word document."""
        for etype, content in elements:
            if etype == 'heading1':
                add_heading(doc, content, level=1)
            elif etype == 'heading2':
                add_heading(doc, content, level=2)
            elif etype == 'heading3':
                para = doc.add_paragraph()
                run = para.add_run(content)
                set_run_font(run, size=Pt(11), bold=True, italic=True)
                set_paragraph_format(para, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                                     space_before=Pt(12), space_after=Pt(4))
            elif etype == 'paragraph':
                add_body_paragraph(doc, content)
            elif etype == 'table_caption':
                para = doc.add_paragraph()
                add_formatted_text(para, content, size=Pt(10))
                set_paragraph_format(para, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                                     space_before=Pt(12), space_after=Pt(4))
            elif etype == 'table':
                header, rows = content
                add_table(doc, header, rows)
                para = doc.add_paragraph()
                set_paragraph_format(para, space_after=Pt(2), space_before=Pt(0))
            elif etype == 'table_note':
                para = doc.add_paragraph()
                text = content.strip('*').strip()
                run = para.add_run(text)
                set_run_font(run, size=Pt(9), italic=True)
                set_paragraph_format(para, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                                     space_before=Pt(2), space_after=Pt(8))
            elif etype == 'list_item':
                para = doc.add_paragraph()
                add_formatted_text(para, content, size=FONT_SIZE)
                set_paragraph_format(para, space_before=Pt(2), space_after=Pt(2))
                para.paragraph_format.left_indent = Inches(0.5)
            elif etype == 'bullet':
                para = doc.add_paragraph()
                add_formatted_text(para, "• " + content, size=FONT_SIZE)
                set_paragraph_format(para, space_before=Pt(2), space_after=Pt(2))
                para.paragraph_format.left_indent = Inches(0.5)

    for fname in sections_files:
        md_text = (DRAFTS / fname).read_text(encoding='utf-8')
        elements = extract_body_and_tables(md_text)
        render_elements(elements)

        # Insert figures after this section
        if fname in figure_after_section:
            for fig_path, fig_caption, fig_width in figure_after_section[fname]:
                if fig_path.exists():
                    add_figure(doc, fig_path, fig_caption, width=fig_width)

    # ==================== METHODS ====================
    methods_md = (DRAFTS / "methods_v3.md").read_text(encoding='utf-8')
    elements = extract_body_and_tables(methods_md)
    render_elements(elements)

    # ==================== REFERENCES (placeholder) ====================
    add_heading(doc, "References", level=1)
    refs = [
        "An, L. Modeling human decisions in coupled human and natural systems: review of agent-based models. Ecol. Model. 229, 25–36 (2012).",
        "Bankes, S. Exploratory modeling for policy analysis. Oper. Res. 41, 435–449 (1993).",
        "Berglund, E. Z. Using agent-based modeling for water resources planning and management. J. Water Resour. Plan. Manage. 141, 04015025 (2015).",
        "Blair, P. & Buytaert, W. Socio-hydrological modelling: a review asking 'why, what and how?'. Hydrol. Earth Syst. Sci. 20, 443–478 (2016).",
        "Bonabeau, E. Agent-based modeling: methods and techniques for simulating human systems. Proc. Natl Acad. Sci. USA 99, 7280–7287 (2002).",
        "Bubeck, P. et al. A review of risk perceptions and other factors that influence flood mitigation behavior. Risk Anal. 32, 1481–1495 (2012).",
        "Castilla-Rho, J. C. et al. An agent-based platform for simulating complex human–aquifer interactions in managed groundwater systems. Environ. Model. Softw. 92, 27–46 (2017).",
        "Di Baldassarre, G. et al. Socio-hydrology: conceptualising human-flood interactions. Hydrol. Earth Syst. Sci. 17, 3295–3303 (2013).",
        "Di Baldassarre, G. et al. Debates: perspectives on socio-hydrology: capturing feedbacks between physical and social processes. Water Resour. Res. 51, 4770–4781 (2015).",
        "Di Baldassarre, G. et al. Sociohydrology: scientific challenges in addressing the sustainable development goals. Water Resour. Res. 55, 6327–6355 (2019).",
        "Epstein, J. M. & Axtell, R. Growing Artificial Societies: Social Science from the Bottom Up (MIT Press, 1996).",
        "Gao, C. et al. Large language models empowered agent-based modeling and simulation: a survey and perspectives. Humanit. Soc. Sci. Commun. 11, 1498 (2024).",
        "Gemma Team. Gemma 3 Technical Report. Preprint at https://arxiv.org/abs/2503.19786 (2025).",
        "Grimm, V. et al. Pattern-oriented modeling of agent-based complex systems: lessons from ecology. Science 310, 987–991 (2005).",
        "Haer, T., Botzen, W. J. W., de Moel, H. & Aerts, J. C. J. H. Integrating household risk mitigation behavior in flood risk analysis: an agent-based model approach. Risk Anal. 37, 1977–1992 (2017).",
        "Huang, L. et al. A survey on hallucination in large language models: principles, taxonomy, challenges, and open questions. ACM Trans. Inf. Syst. 43, 1–55 (2025).",
        "Hung, F. & Yang, Y.-C. E. Assessing adaptive irrigation impacts on water scarcity in nonstationary environments — a multi-agent reinforcement learning approach. Water Resour. Res. 57, e2020WR029262 (2021).",
        "Hung, F. & Yang, Y.-C. E. Assessing the impact of adaptation strategies on water scarcity under climate change in the Colorado River Basin. J. Hydrol. 612, 128193 (2022).",
        "Hyun, J.-Y. & Yang, Y.-C. E. Using a coupled agent-based modeling approach to analyze the role of risk perception in water management decisions. Hydrol. Earth Syst. Sci. 23, 2261–2278 (2019).",
        "Lin, C.-Y. & Yang, Y.-C. E. An investigation of coupled natural human systems using a two-way coupled agent-based modeling framework. Environ. Model. Softw. 155, 105451 (2022).",
        "Liu, J. et al. Complexity of coupled human and natural systems. Science 317, 1513–1516 (2007).",
        "Loucks, D. P. & van Beek, E. Water Resource Systems Planning and Management: An Introduction to Methods, Models, and Applications (Springer, 2017).",
        "Maass, A. et al. Design of Water-Resource Systems: New Techniques for Relating Economic Objectives, Engineering Analysis, and Governmental Planning (Harvard Univ. Press, 1962).",
        "Müller, B. et al. Describing human decisions in agent-based models — ODD+D, an extension of the ODD protocol. Environ. Model. Softw. 48, 37–48 (2013).",
        "Ostrom, E. Governing the Commons: The Evolution of Institutions for Collective Action (Cambridge Univ. Press, 1990).",
        "Park, J. S. et al. Generative agents: interactive simulacra of human behavior. In Proc. 36th ACM Symp. on User Interface Software and Technology 1–22 (2023).",
        "Rogers, R. W. Cognitive and physiological processes in fear appeals and attitude change: a revised theory of protection motivation. In Social Psychophysiology: A Sourcebook (eds Cacioppo, J. T. & Petty, R. E.) 153–176 (Guilford, 1983).",
        "Schlüter, M. et al. A framework for mapping and comparing behavioural theories in models of social-ecological systems. Ecol. Econ. 131, 21–35 (2017).",
        "Sivapalan, M. et al. Socio-hydrology: a new science of people and water. Hydrol. Process. 26, 1270–1276 (2012).",
        "Vezhnevets, A. S. et al. Generative agent-based modeling with actions grounded in physical, social, or digital space using Concordia. Preprint at https://arxiv.org/abs/2312.03664 (2023).",
        "Yang, Y.-C. E., Cai, X. & Stipanović, D. M. A decentralized optimization algorithm for multiagent system-based watershed management. Water Resour. Res. 45, W08430 (2009).",
    ]
    for ref in refs:
        para = doc.add_paragraph()
        add_formatted_text(para, ref, size=Pt(10))
        set_paragraph_format(para, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                             space_before=Pt(1), space_after=Pt(1))
        para.paragraph_format.left_indent = Inches(0.3)
        para.paragraph_format.first_line_indent = Inches(-0.3)

    # ==================== END MATTER ====================
    endmatter_sections = [
        ("Data Availability",
         "All simulation output data, agent decision traces, and governance audit logs "
         "generated in this study will be deposited in Zenodo upon acceptance (DOI to be assigned). "
         "Raw CRSS demand data for the irrigation domain are publicly available from the US "
         "Bureau of Reclamation (https://www.usbr.gov/lc/region/g4000/NaturalFlow/). Flood "
         "depth grids were generated from publicly available FEMA flood maps. Census-tract "
         "demographic profiles were derived from the US Census Bureau American Community Survey."),
        ("Code Availability",
         "The Water Agent Governance Framework (WAGF) source code, including the broker "
         "architecture, governance validators, configuration files for both domains, and all "
         "analysis scripts used to generate the figures and tables in this paper, will be "
         "archived on Zenodo and made available on GitHub upon acceptance (DOI and URL to be assigned). "
         "Experiments were run using Ollama (v0.5.x) with open-weight models (Gemma-3 and Ministral "
         "families); no proprietary APIs were required."),
        ("Author Contributions",
         "W.-Y.C.: Conceptualization, Methodology, Software, Validation, Formal Analysis, "
         "Investigation, Data Curation, Writing \u2014 Original Draft, Visualization. "
         "Y.C.E.Y.: Supervision, Writing \u2014 Review & Editing, Funding Acquisition."),
        ("Competing Interests",
         "The authors declare no competing interests."),
        ("Acknowledgements",
         "[To be added \u2014 funding sources, computational resources, etc.]"),
    ]
    for title, body in endmatter_sections:
        add_heading(doc, title, level=2)
        para = doc.add_paragraph()
        add_formatted_text(para, body, size=FONT_SIZE)
        set_paragraph_format(para, first_line_indent=None)

    # Line numbers
    add_line_numbers(doc)

    # Page numbers in footer
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        set_paragraph_format(para, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                             line_spacing=WD_LINE_SPACING.SINGLE)
        # Add page number field
        run = para.add_run()
        fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run._element.append(fldChar1)
        run2 = para.add_run()
        instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run2._element.append(instrText)
        run3 = para.add_run()
        fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run3._element.append(fldChar2)

    out_path = OUTPUT / "NatureWater_MainText_v16.docx"
    doc.save(str(out_path))
    print(f"Main paper saved to: {out_path}")
    return out_path


def compile_si():
    """Compile the Supplementary Information as a separate document."""
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    # Default font
    style = doc.styles['Normal']
    style.font.name = FONT_NAME
    style.font.size = FONT_SIZE
    style.paragraph_format.line_spacing_rule = LINE_SPACING
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Title
    para = doc.add_paragraph()
    run = para.add_run("Supplementary Information")
    set_run_font(run, size=Pt(16), bold=True)
    set_paragraph_format(para, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                         space_before=Pt(24), space_after=Pt(6))

    para = doc.add_paragraph()
    run = para.add_run("Institutional Constraints Widen Adaptive Behavioural Diversity in Language-Based Water Agents")
    set_run_font(run, size=Pt(12), italic=True)
    set_paragraph_format(para, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                         space_after=Pt(24))

    # Parse SI
    si_md = (DRAFTS / "supplementary_information.md").read_text(encoding='utf-8')
    elements = extract_body_and_tables(si_md)

    for etype, content in elements:
        if etype == 'heading1':
            add_heading(doc, content, level=1)
        elif etype == 'heading2':
            add_heading(doc, content, level=2)
        elif etype == 'heading3':
            para = doc.add_paragraph()
            run = para.add_run(content)
            set_run_font(run, size=Pt(11), bold=True, italic=True)
            set_paragraph_format(para, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                                 space_before=Pt(12), space_after=Pt(4))
        elif etype == 'paragraph':
            add_body_paragraph(doc, content)
        elif etype == 'table_caption':
            para = doc.add_paragraph()
            add_formatted_text(para, content, size=Pt(10))
            set_paragraph_format(para, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                                 space_before=Pt(12), space_after=Pt(4))
        elif etype == 'table':
            header, rows = content
            add_table(doc, header, rows)
        elif etype == 'table_note':
            para = doc.add_paragraph()
            text = content.strip('*').strip()
            run = para.add_run(text)
            set_run_font(run, size=Pt(9), italic=True)
            set_paragraph_format(para, space_before=Pt(2), space_after=Pt(8))
        elif etype == 'list_item':
            para = doc.add_paragraph()
            add_formatted_text(para, content, size=FONT_SIZE)
            set_paragraph_format(para, space_before=Pt(2), space_after=Pt(2))
            para.paragraph_format.left_indent = Inches(0.5)
        elif etype == 'bullet':
            para = doc.add_paragraph()
            add_formatted_text(para, "• " + content, size=FONT_SIZE)
            set_paragraph_format(para, space_before=Pt(2), space_after=Pt(2))
            para.paragraph_format.left_indent = Inches(0.5)

    # Line numbers
    add_line_numbers(doc)

    out_path = OUTPUT / "NatureWater_SI_v16.docx"
    doc.save(str(out_path))
    print(f"SI saved to: {out_path}")
    return out_path


if __name__ == "__main__":
    print("Compiling Nature Water paper v16...")
    main_path = compile_main_paper()
    si_path = compile_si()
    print(f"\nDone. Files:")
    print(f"  Main: {main_path}")
    print(f"  SI:   {si_path}")
