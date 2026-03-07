#!/usr/bin/env python3
"""Generate Word document from 10-minute talking points markdown."""

import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import re

BLUE = RGBColor(0x2F, 0x54, 0x96)
GREY = RGBColor(0x80, 0x80, 0x80)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BLUE = "D6E4F0"
FONT_NAME = "Calibri"
BODY_PT = 10


def set_run_font(run, size=BODY_PT, bold=False, italic=False, color=None, name=FONT_NAME):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    # Set East Asian font
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
        r.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), name)


def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    tc = cell._element
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = parse_xml(f'<w:tcPr {nsdecls("w")}></w:tcPr>')
        tc.insert(0, tcPr)
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
    )
    tcPr.append(shading)


def add_styled_paragraph(doc, text, style_name=None, bold=False, italic=False,
                         color=None, size=BODY_PT, indent_left=None,
                         space_before=0, space_after=4):
    """Add a paragraph with consistent font styling."""
    if style_name:
        p = doc.add_paragraph(style=style_name)
    else:
        p = doc.add_paragraph()

    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if indent_left is not None:
        pf.left_indent = Inches(indent_left)

    # Parse bold markers within text
    add_rich_text(p, text, default_bold=bold, default_italic=italic,
                  default_color=color, default_size=size)
    return p


def add_rich_text(paragraph, text, default_bold=False, default_italic=False,
                  default_color=None, default_size=BODY_PT):
    """Parse **bold** markers in text and add runs accordingly."""
    # Split on **...**
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            inner = part[2:-2]
            run = paragraph.add_run(inner)
            set_run_font(run, size=default_size, bold=True,
                         italic=default_italic, color=default_color)
        else:
            if part:
                run = paragraph.add_run(part)
                set_run_font(run, size=default_size, bold=default_bold,
                             italic=default_italic, color=default_color)


def add_heading1(doc, text):
    """Add a blue Heading 1."""
    p = doc.add_heading(level=1)
    run = p.add_run(text)
    set_run_font(run, size=14, bold=True, color=BLUE)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_quote(doc, text):
    """Add an italic indented quote paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    # Strip leading/trailing quotes
    clean = text.strip()
    if clean.startswith('"') and clean.endswith('"'):
        pass  # keep the quotes
    add_rich_text(p, clean, default_italic=True, default_size=BODY_PT)
    return p


def add_bullet(doc, text):
    """Add a bullet point with rich text support."""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    add_rich_text(p, text, default_size=BODY_PT)
    return p


def add_numbered(doc, text, restart=False):
    """Add a numbered list item."""
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    add_rich_text(p, text, default_size=BODY_PT)
    return p


def add_table(doc, header_row, data_rows):
    """Add a formatted table with blue header and alternating shading."""
    n_cols = len(header_row)
    table = doc.add_table(rows=1 + len(data_rows), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    hdr = table.rows[0]
    for i, txt in enumerate(header_row):
        cell = hdr.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(txt)
        set_run_font(run, size=9, bold=True, color=WHITE)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "2F5496")

    # Data rows
    for r_idx, row_data in enumerate(data_rows):
        row = table.rows[r_idx + 1]
        for c_idx, txt in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            add_rich_text(p, txt, default_size=9)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if r_idx % 2 == 1:
                set_cell_shading(cell, LIGHT_BLUE)

    return table


def main():
    brief_dir = Path(__file__).resolve().parents[1] / "professor_briefing"
    src = brief_dir / "10min_talking_points.md"
    dst = brief_dir / "10min_talking_points.docx"

    with open(src, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = FONT_NAME
    style.font.size = Pt(BODY_PT)
    style.paragraph_format.space_after = Pt(4)

    # Set List Bullet font
    for sname in ['List Bullet', 'List Number']:
        if sname in doc.styles:
            s = doc.styles[sname]
            s.font.name = FONT_NAME
            s.font.size = Pt(BODY_PT)

    # ---- TITLE ----
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(2)
    run = title_p.add_run("WAGF Framework \u2014 10-Minute Presentation Talking Points")
    set_run_font(run, size=16, bold=True, color=BLUE)

    # ---- SUBTITLE ----
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_after = Pt(4)
    run = sub_p.add_run("For: CS Professor Committee Meeting")
    set_run_font(run, size=10, color=GREY)

    sub2_p = doc.add_paragraph()
    sub2_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub2_p.paragraph_format.space_after = Pt(12)
    run = sub2_p.add_run("Format: Framework explanation + preliminary results")
    set_run_font(run, size=10, color=GREY)

    # ---- Process content ----
    i = 5  # skip title / subtitle / format lines
    while i < len(lines):
        line = lines[i].rstrip('\n')
        stripped = line.strip()

        # Skip horizontal rules
        if stripped == '---':
            i += 1
            continue

        # Heading 2 → Heading 1 (slide sections)
        if stripped.startswith('## '):
            heading_text = stripped[3:].strip()
            # Skip the "For:" and "Format:" subtitles we already added
            if heading_text.startswith('For:') or heading_text.startswith('Format:'):
                i += 1
                continue
            add_heading1(doc, heading_text)
            i += 1
            continue

        # Blockquote
        if stripped.startswith('> '):
            quote_text = stripped[2:].strip()
            add_quote(doc, quote_text)
            i += 1
            continue

        # Table detection
        if stripped.startswith('|') and i + 1 < len(lines):
            # Collect table lines
            table_lines = []
            j = i
            while j < len(lines) and lines[j].strip().startswith('|'):
                table_lines.append(lines[j].strip())
                j += 1

            if len(table_lines) >= 3:
                # Parse header
                header = [c.strip() for c in table_lines[0].split('|')[1:-1]]
                # Skip separator row (index 1)
                data = []
                for tl in table_lines[2:]:
                    row = [c.strip() for c in tl.split('|')[1:-1]]
                    data.append(row)
                add_table(doc, header, data)

            i = j
            continue

        # Bullet point (- prefix)
        if stripped.startswith('- '):
            bullet_text = stripped[2:].strip()
            add_bullet(doc, bullet_text)
            i += 1
            continue

        # Numbered list (1. 2. 3. etc.)
        m = re.match(r'^(\d+)\.\s+(.*)', stripped)
        if m:
            item_text = m.group(2).strip()
            # Check for sub-bullets
            add_numbered(doc, item_text)
            # Consume sub-items indented under this number
            i += 1
            while i < len(lines):
                sub = lines[i].rstrip('\n')
                sub_stripped = sub.strip()
                if sub_stripped.startswith('- ') and (sub.startswith('   ') or sub.startswith('\t')):
                    sub_text = sub_stripped[2:].strip()
                    p = doc.add_paragraph(style='List Bullet 2')
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(1)
                    add_rich_text(p, sub_text, default_size=BODY_PT)
                    i += 1
                else:
                    break
            continue

        # Bold label lines (e.g., "**Opening hook:**")
        if stripped.startswith('**') and stripped.endswith('**'):
            label = stripped[2:-2]
            add_styled_paragraph(doc, label, bold=True, space_before=6, space_after=2)
            i += 1
            continue

        # Bold label with trailing text (e.g., "**Key points to hit:**")
        if stripped.startswith('**') and '**' in stripped[2:]:
            # It's a bold-prefixed line
            add_styled_paragraph(doc, stripped, space_before=6, space_after=2)
            i += 1
            continue

        # Plain text (non-empty)
        if stripped:
            add_styled_paragraph(doc, stripped, space_before=2, space_after=2)

        i += 1

    # Save
    doc.save(dst)
    print(f"Word document saved to: {dst}")


if __name__ == '__main__':
    main()
