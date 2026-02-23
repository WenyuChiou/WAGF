"""Compile Nature Water v14 sections into NW_Full_v14.docx"""
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

# -- Style setup --
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(11)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(6)

# Heading styles
for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Times New Roman'
    hs.font.color.rgb = RGBColor(0, 0, 0)
    hs.font.bold = True
    if level == 1:
        hs.font.size = Pt(14)
    elif level == 2:
        hs.font.size = Pt(12)
    else:
        hs.font.size = Pt(11)

def add_paragraph(text, bold=False, italic=False, style_name='Normal', alignment=None):
    p = doc.add_paragraph(style=style_name)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if alignment:
        p.alignment = alignment
    return p

def add_table(header_line, data_lines, caption=None, footnote=None):
    """Parse markdown table lines into a Word table."""
    if caption:
        p = doc.add_paragraph()
        run = p.add_run(caption)
        run.bold = True
        run.font.size = Pt(10)

    # Parse headers
    headers = [c.strip().replace('**', '') for c in header_line.split('|') if c.strip()]

    ncols = len(headers)
    nrows = len(data_lines) + 1

    table = doc.add_table(rows=nrows, cols=ncols)
    table.style = 'Table Grid'
    table.autofit = True

    # Header row
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = 'Times New Roman'

    # Data rows
    for i, line in enumerate(data_lines):
        cells = [c.strip().replace('**', '') for c in line.split('|') if c.strip()]
        for j, val in enumerate(cells):
            if j < ncols:
                cell = table.rows[i+1].cells[j]
                cell.text = ''
                run = cell.paragraphs[0].add_run(val)
                run.font.size = Pt(9)
                run.font.name = 'Times New Roman'

    if footnote:
        p = doc.add_paragraph()
        run = p.add_run(footnote)
        run.italic = True
        run.font.size = Pt(9)

# =====================================================================
# TITLE
# =====================================================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Institutional Governance Enables Adaptive Strategy Diversity\nin Language-Based Water Resource Simulation')
run.bold = True
run.font.size = Pt(16)
run.font.name = 'Times New Roman'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('[Authors]')
run.font.size = Pt(11)

doc.add_paragraph()  # spacer

# =====================================================================
# ABSTRACT
# =====================================================================
doc.add_heading('Abstract', level=1)
abstract_text = (
    'We propose a governance architecture that validates natural-language agent reasoning '
    'against modular physical and institutional rules. In a 42-year Colorado River irrigation '
    'simulation, governed agents extracted more water than ungoverned agents while maintaining '
    'drought responsiveness \u2014 adaptive exploitation consistent with prior-appropriation dynamics '
    'that parameterized models cannot represent. Institutional rules created feasibility boundaries '
    'within which agents developed differentiated strategies autonomously, producing higher strategy '
    'diversity than both ungoverned agents and a hand-coded baseline. Targeted removal of a single '
    'rule \u2014 the demand ceiling linking individual proposals to basin-wide demand \u2014 increased '
    'diversity but decoupled it from drought signals, nearly doubling shortage frequency and '
    'distinguishing adaptive from arbitrary diversity. The effect generalized from chronic drought '
    '(78 agents, 42 years) to acute flood hazard (100 agents, 10 years), was positive across six '
    'model scales tested, and significant for five. Institutional boundaries widen adaptive capacity '
    'in ways that fixed decision rules cannot represent.'
)
add_paragraph(abstract_text)
doc.add_page_break()

# =====================================================================
# INTRODUCTION
# =====================================================================
doc.add_heading('Introduction', level=1)

intro_paragraphs = [
    'For six decades, computational models of water resource systems have represented human decisions in a single paradigm: as computable mappings from system states to actions. Whether the mapping is a linear program, a coupled differential equation, or an agent\u2019s decision rule, the person choosing how much water to divert, whether to buy flood insurance, or when to reduce irrigation is reduced to a parameterized function. This paradigm \u2014 established when the Harvard Water Program introduced operations research to water planning (Maass et al., 1962) \u2014 enabled rigorous engineering analysis but compressed the reasoning process into a mathematical object, foreclosing a question that water governance increasingly demands we answer: not just what people decide, but how they reason toward those decisions.',

    'Subsequent frameworks expanded the scope of human\u2013water modelling without departing from this numerical paradigm. Coupled human\u2013natural systems theory recognized that human decisions and physical processes co-evolve through feedback loops (Liu et al., 2007). Sociohydrology formalized water-specific feedbacks \u2014 levee effects, supply\u2013demand cycles, reservoir operation \u2014 as coupled differential equations, but human behaviour entered these equations as aggregate parameters rather than as individual reasoning processes (Sivapalan et al., 2012; Di Baldassarre et al., 2019). As Blair and Buytaert (2016) observed in their review of socio-hydrological modelling, models could represent that populations respond to flood risk, but not how individuals reason toward protective action.',

    'Agent-based modelling introduced individual-level decision-making by representing heterogeneous agents who act on local information and interact through shared environments (Epstein and Axtell, 1996; Bonabeau, 2002). In water research, Berglund (2015) modelled irrigation scheduling in the Yakima Basin through threshold-based allocation rules; Haer et al. (2017) used Protection Motivation Theory to drive household flood-adaptation choices through threat and coping appraisal scores; Hung and Yang (2021) encoded prior-appropriation operating rules for Colorado River demand management. More recent approaches have introduced adaptive mechanisms \u2014 Bayesian belief updating, reinforcement learning policies (Castilla-Rho et al., 2017), and BDI cognitive architectures \u2014 that allow agents to learn and extrapolate beyond initial conditions. Yet even these approaches map numerical state vectors to numerical actions; the decision architecture \u2014 state representation, action space, and behavioural theory \u2014 must be redesigned for each domain (An, 2012; M\u00fcller et al., 2013), drawing from a fragmented landscape of over 40 distinct behavioural theories that Schl\u00fcter et al. (2017) mapped across agent-based applications. Meanwhile, all agents within a given simulation follow the same cognitive model, differing only in parameter values \u2014 a structural rigidity that limits the representation of qualitatively different reasoning strategies within a single population.',

    'Large language models introduce a qualitatively different representational format: agents that reason in natural language rather than through numerical functions. Instead of mapping state variables to actions through parameterized functions, a language-based agent reads contextual information \u2014 drought indices, water rights, neighbour behaviour, institutional announcements \u2014 and generates a decision through linguistic reasoning that can reference domain knowledge, weigh trade-offs, and articulate justifications. Generative agents have demonstrated believable social behaviours including persona maintenance and adaptive planning (Park et al., 2023); structured environments such as Concordia have shown that language agents can operate within defined action spaces (Vezhnevets et al., 2023); and large-scale social simulations have begun exploring population-level dynamics with over 1,000 LLM agents (Gao et al., 2024). This body of work suggests that natural language may serve as a medium for representing the reasoning processes that numerical formats compress away.',

    'Yet language-based agents carry risks that are acute for scientific simulation. An agent might propose a water allocation that violates mass balance, a home elevation on an already-elevated structure, or a demand increase beyond its legal water right. These constraint violations \u2014 distinct from the factual hallucinations studied in natural language processing (Huang et al., 2025) \u2014 arise because LLMs lack inherent physical grounding, and their frequency varies substantially across model families and scales. Most existing LLM-agent studies lack empirical behavioural benchmarking against observed patterns. The question is not whether language-based agents can reason about water decisions \u2014 early evidence suggests they can \u2014 but whether their outputs can be governed by the physical and institutional constraints that define real water systems.',

    'Answering this governance question connects computation to institutional theory. Constraints are conventionally expected to reduce the space of available actions. But Ostrom (1990) observed that well-designed institutions for managing common-pool resources do not merely restrict behaviour \u2014 they define feasibility boundaries within which diverse adaptive strategies become viable. We hypothesized that architectural governance \u2014 validating agent proposals against physical and institutional rules before execution \u2014 would produce a structurally analogous outcome at the computational level: eliminating physically impossible outputs while preserving or expanding the space of plausible decisions. We note that this is a structural parallel between institutional rules governing communities and computational validators governing artificial agents, requiring empirical rather than theoretical justification.',

    'Here we test this hypothesis across two contrasting water domains: irrigation management in the Colorado River Basin (78 agents, 42 years, generating over 9,800 governed decisions across three seeds) and household flood adaptation (100 agents, 10 years, stochastic flood events, over 5,000 governed decisions validated against empirical behavioural benchmarks). Governed agents exploit water more aggressively during abundance and curtail during drought \u2014 a pattern of adaptive exploitation that neither mathematical optimization nor rule-based agents can represent, and that requires institutional rules to emerge. Ungoverned agents collapse into repetitive demand patterns. This governance effect is positive for all six model scales tested (3B to 27B parameters, two model families), statistically significant for five, and persists among first-attempt proposals before any governance feedback. Because each institutional rule can be independently enabled, disabled, or reconfigured, the approach functions as a method for experimentally probing how specific institutional designs shape adaptive water behaviour \u2014 a computationally governed representation of how people reason about water under uncertainty.',
]

for para in intro_paragraphs:
    add_paragraph(para)

doc.add_page_break()

# =====================================================================
# RESULTS
# =====================================================================
doc.add_heading('Results', level=1)

# --- R1 ---
doc.add_heading('Institutional rules enable adaptive exploitation of water resources', level=2)

r1_paras = [
    'Governed agents extracted more water than ungoverned agents while maintaining stronger coupling between individual decisions and reservoir state (Table 1). Over 42 simulated years, governed agents achieved a mean demand ratio of 0.394 compared with 0.288 for ungoverned agents \u2014 yet governed agents responded more sensitively to drought, as reflected in the correlation between annual Lake Mead elevation and aggregate demand (r = 0.547 governed versus 0.378 ungoverned). Lake Mead was consequently lower under governance (42-year mean 1,094 ft versus 1,173 ft) and governed agents triggered shortage conditions more frequently (13.3 versus 5.0 of 42 years). The minimum Mead elevation was comparable (governed 1,002 ft versus ungoverned 1,001 ft), confirming that both systems reached the physical floor during severe early drought; the divergence occurred during recovery, where governed agents adaptively adjusted demand while ungoverned agents could not.',

    'This pattern \u2014 higher extraction during abundance, proportionate response to scarcity signals \u2014 is consistent with adaptive exploitation under prior-appropriation. In real Colorado River management, institutional rules (shortage tiers, delivery obligations) enable senior rights holders to utilize water aggressively in normal years because those rules guarantee curtailment during drought. Governed agents display a structurally analogous dynamic: institutional validators made it safe to extract more water by ensuring that drought-inappropriate proposals were blocked before execution. The water-system consequence is that governance shifts the operating point upward without degrading drought responsiveness.',

    'Without governance, agents collapsed into monotonic demand-increase patterns, concentrating 77\u201382% of decisions on demand increases across all seeds. Their demand trajectories converged toward a code-level ceiling over 42 years, producing artificially conservative demand profiles driven by clamps rather than adaptive choice. Ungoverned reservoir stability was inertial \u2014 agents rarely triggered shortage because they never extracted enough to draw Mead below critical thresholds, not because they responded to drought. This adaptive exploitation pattern \u2014 visible because language-based agents generate explicit reasoning that governance can evaluate and channel \u2014 represents a water-system dynamic that conventional parameterized decision rules cannot produce.',
]
for para in r1_paras:
    add_paragraph(para)

# Table 1
add_table(
    '| Metric | Governed | Ungoverned | A1 (No Ceiling) |',
    [
        '| Mean demand ratio | 0.394 \u00b1 0.004 | 0.288 \u00b1 0.020 | 0.440 \u00b1 0.012 |',
        '| 42-yr mean Mead elevation (ft) | 1,094 | 1,173 | 1,069 |',
        '| Demand\u2013Mead coupling (r) | 0.547 \u00b1 0.083 | 0.378 \u00b1 0.081 | 0.234 \u00b1 0.127 |',
        '| Shortage years (/42) | 13.3 \u00b1 1.5 | 5.0 \u00b1 1.7 | 25.3 \u00b1 1.5 |',
        '| Min Mead elevation (ft) | 1,002 \u00b1 1 | 1,001 \u00b1 0.4 | 984 \u00b1 11 |',
        '| Strategy diversity (EHE) | 0.738 \u00b1 0.017 | 0.637 \u00b1 0.017 | 0.793 \u00b1 0.002 |',
        '| Behavioural Rationality (BRI, %) | 58.0 | 9.4 | \u2014 |',
    ],
    caption='Table 1. Water-system outcomes and strategy diversity across three governance conditions (irrigation domain, Gemma-3 4B, 78 agents \u00d7 42 years, 3 runs each).',
    footnote='Three independent runs per condition (seeds 42, 43, 44). Demand ratio = requested volume / historical baseline allocation. Demand\u2013Mead coupling = Pearson r between annual Lake Mead elevation and annual mean demand ratio (positive r indicates agents reduce demand during drought). EHE = Effective Heterogeneity Entropy (normalized Shannon entropy over 5 action types; see Methods). BRI = fraction of high-scarcity decisions where agents did not increase demand (null expectation under uniform random = 60%). A1 removes the demand ceiling stabilizer only (see next section). See Supplementary Table S6 for additional water-system metrics.'
)

doc.add_paragraph()  # spacer

# --- R2 ---
doc.add_heading('A single institutional rule couples individual decisions to basin-wide drought', level=2)

r2_paras = [
    'To identify which institutional rules create the coupling between individual decisions and reservoir state, we removed a single validator \u2014 the demand ceiling stabilizer, which blocks demand-increase proposals when aggregate basin demand exceeds 6.0 MAF \u2014 while retaining all eleven other validators (condition A1; see Methods).',

    'Removing this one rule of twelve collapsed demand\u2013Mead coupling from r = 0.547 to 0.234, nearly doubled shortage years from 13.3 to 25.3, and dropped minimum Mead elevation to 984 ft below the Tier 3 shortage threshold (Table 1, A1 column). Yet removing the ceiling increased strategy diversity: EHE rose from 0.738 to 0.793. Agents diversified further \u2014 but into extraction patterns decoupled from drought signals, producing higher demand ratios (0.440 versus 0.394) with weaker environmental responsiveness.',

    'This establishes a distinction central to interpreting the governance effect: between diversity (a wider action distribution) and adaptive diversity (an action distribution coupled to environmental state). The demand ceiling does not suppress diversity; it channels diversity toward drought-responsive patterns. Without it, agents diversify into individually rational but collectively maladaptive extraction \u2014 the hallmark of commons dilemmas that Ostrom (1990) identified as the target of institutional design.',

    'The demand ceiling is the only one of twelve validators linking individual proposals to aggregate basin state. Its removal demonstrates that governance-induced diversity is functionally adaptive, not a statistical artefact of constraint-based rejection. This experimental decomposition \u2014 isolating one institutional rule\u2019s contribution to human\u2013water coupling \u2014 would be impossible in conventional agent-based models where institutional rules are embedded in code and cannot be independently manipulated.',
]
for para in r2_paras:
    add_paragraph(para)

# --- R3 ---
doc.add_heading('Governance generates strategy diversity beyond what hand-coded models can represent', level=2)

r3_paras = [
    'Governed agents exhibited higher strategy diversity than both ungoverned agents and a hand-coded Protection Motivation Theory baseline across both water domains. In irrigation, governed EHE (0.738 \u00b1 0.017) exceeded ungoverned (0.637 \u00b1 0.017) with zero distributional overlap across three seeds (Table 1). In flood adaptation, the ordering was consistent: governed language agents (0.752 \u00b1 0.052) exceeded rule-based PMT agents (0.689 \u00b1 0.001), which exceeded ungoverned language agents (0.337 \u00b1 0.064; Table 2). Ungoverned agents collapsed into behavioural monoculture: 77\u201382% demand increases in irrigation, 85.9% inaction in flood.',

    'This diversity is generated by agents reasoning within governance, not filtered into existence by validators. Proposals submitted before any governance feedback already showed higher diversity (first-attempt EHE 0.761 governed versus 0.640 ungoverned; irrigation domain), confirming that the governance context shapes the reasoning process rather than the rejection-retry mechanism creating diversity post hoc (see Supplementary Information for retry statistics).',

    'The rule-based PMT agent\u2019s diversity stems entirely from parameterized variation: agents differ in income, flood zone, and prior experience, producing different threshold crossings, but all follow identical decision logic. Governed language agents achieve higher diversity through qualitatively different reasoning paths \u2014 agents develop rationales referencing personal trade-offs, contextual factors, and institutional constraints in ways that threshold-based rules cannot represent (see Supplementary Information for paired reasoning traces). This is not a marginal improvement in the same representational format; it is a shift from parameterized homogeneity to reasoning-generated heterogeneity.',
]
for para in r3_paras:
    add_paragraph(para)

# Table 2
add_table(
    '| Condition | EHE | CACR (%) | do_nothing (%) | insurance (%) | elevation (%) | relocation (%) |',
    [
        '| Governed LLM | 0.752 \u00b1 0.052 | 100.0 | 35.6 | 50.7 | 10.6 | 3.0 |',
        '| Rule-based PMT | 0.689 \u00b1 0.001 | 100.0 | 10.6 | 49.1 | 40.2 | 0.1 |',
        '| Ungoverned LLM | 0.337 \u00b1 0.064 | 85.5 | 85.9 | 11.7 | 2.3 | 0.0 |',
    ],
    caption='Table 2. Strategy diversity: governed LLM vs rule-based PMT vs ungoverned LLM (flood domain, Gemma-3 4B, 100 agents \u00d7 10 years, 3 runs each).',
    footnote='EHE computed from annual action selections. CACR = Construct-Action Coherence Rate: fraction of decisions where the agent\u2019s stated risk assessment is consistent with its chosen action (see Methods; operationalized differently from the irrigation-domain BRI in Table 1). Rule-based agent uses deterministic PMT threshold logic with parameterized agent heterogeneity; composite recommendations (simultaneous insurance + elevation) split into constituents for EHE computation (see Methods).'
)

doc.add_paragraph()

# --- R4 ---
doc.add_heading('The governance effect generalizes across water hazard types and model architectures', level=2)

r4_paras = [
    'The governance mechanism that produces adaptive exploitation in chronic drought also generates higher strategy diversity under acute flood hazard. In the flood domain, 100 household agents made protective decisions (insurance, elevation, relocation, or inaction) over 10 years with stochastic flood events \u2014 a fundamentally different water context from continuous irrigation allocation. Governance eliminated reasoning-action mismatches entirely (0.0% across all models versus 0.1\u201362.3% ungoverned).',

    'Six models spanning two families and three parameter scales all showed positive governance effects on strategy diversity, five statistically significant (Table 3). The effect was strongest where ungoverned agents exhibited behavioural monoculture: Gemma-3 4B (delta = +0.415) and Ministral 3B (+0.302). Only Gemma-3 12B produced a non-significant effect (+0.012), reflecting strong instruction-tuned priors at this scale (see SI for discussion).',
]
for para in r4_paras:
    add_paragraph(para)

# Table 3
add_table(
    '| Model | Ungoverned EHE | Governed EHE | Delta | 95% CI |',
    [
        '| Gemma-3 4B | 0.337 \u00b1 0.064 | 0.752 \u00b1 0.052 | +0.415 | [+0.393, +0.458] |',
        '| Gemma-3 12B | 0.471 \u00b1 0.014 | 0.483 \u00b1 0.042 | +0.012 | [\u20130.018, +0.039] |',
        '| Gemma-3 27B | 0.462 \u00b1 0.032 | 0.676 \u00b1 0.018 | +0.214 | [+0.204, +0.231] |',
        '| Ministral 3B | 0.431 \u00b1 0.056 | 0.734 \u00b1 0.020 | +0.302 | [+0.232, +0.350] |',
        '| Ministral 8B | 0.579 \u00b1 0.014 | 0.626 \u00b1 0.008 | +0.047 | [+0.042, +0.091] |',
        '| Ministral 14B | 0.665 \u00b1 0.010 | 0.708 \u00b1 0.012 | +0.043 | [+0.041, +0.054] |',
    ],
    caption='Table 3. Governance effect on strategy diversity across six language models (flood domain, 100 agents \u00d7 10 years, 3 runs per condition).',
    footnote='Delta = governed minus ungoverned EHE. 95% CIs from bootstrap resampling (agent-timestep level, 10,000 iterations). Reasoning-action mismatch rates in Table S1.'
)

doc.add_paragraph()

r4_closing = 'The consistency across two water domains, six model scales, two model families, and against a hand-coded baseline provides converging evidence that institutional governance enables adaptive strategy diversity \u2014 a capacity that depends on the governance architecture rather than the specific language model. Because both domains use the same broker with domain-specific rule configurations rather than redesigned decision logic, the architecture addresses a longstanding limitation of agent-based water models: the need to rebuild decision modules for each application domain.'
add_paragraph(r4_closing)

doc.add_page_break()

# =====================================================================
# DISCUSSION
# =====================================================================
doc.add_heading('Discussion', level=1)

discussion_paras = [
    'Governed language-based agents exploit water resources more aggressively during favourable conditions and curtail proportionately during drought \u2014 a pattern of adaptive exploitation that emerges from institutional boundary enforcement rather than programmed decision rules. This addresses the central question of whether computational governance can create conditions for adaptive water management: not by prescribing behaviour, but by defining institutional boundaries within which diverse strategies become viable. Governed agents adjusted demand in response to reservoir state (r = 0.547) substantially more than ungoverned agents (r = 0.378) while extracting a larger share of their water entitlements. In prior-appropriation systems, this parallels how senior rights holders extract their full entitlement in normal years because shortage-sharing rules ensure proportionate curtailment during drought \u2014 the institution does not reduce water use, it makes higher use safe.',

    'The demand ceiling ablation reveals the mechanism. Removing this single rule \u2014 which links individual demand proposals to aggregate basin state \u2014 increased diversity but collapsed demand\u2013Mead coupling and nearly doubled shortage frequency. This separates two properties that institutional theory has long distinguished but never quantified in simulation: diversity (wider repertoire of actions used) and adaptive diversity (actions deployed in response to system state). Ostrom (1990) observed that well-designed institutions for common-pool resource management create structured arenas within which diverse adaptive strategies become viable. Our computational results produce a structurally analogous outcome: governance validators define feasibility boundaries within which language models generate a wider range of valid strategies coupled to environmental conditions. Without these boundaries, agents collapse into repetitive patterns \u2014 not because they lack capacity, but because unconstrained generation produces invalid proposals that crowd out meaningful variation. We emphasize that the parallel is structural: Ostrom\u2019s analysis involves strategic interaction and norm internalization, whereas our validators operate as stateless constraint checks. Whether informal norms emerge from repeated agent interaction remains untested.',

    'The capacity to experimentally decompose institutional rules transforms language-based agents from simulation tools into computational laboratories for water governance. In the ablation demonstrated here, removing one of twelve rules isolated the mechanism coupling individual decisions to basin state. The same approach could test variations in shortage-sharing thresholds (for example, shifting Tier 1 from 1,075 ft to 1,050 ft), alternative allocation regimes (proportional versus priority-based curtailment), premium structures for flood insurance, or drought contingency plan triggers \u2014 observing how populations of reasoning agents endogenously adapt their strategies within physically valid bounds. This capacity is distinct from exploratory modelling (Bankes, 1993) and sensitivity analysis in that the behavioural responses are generated, not sampled from predefined parameter spaces. The rule-based PMT comparison quantifies the representational gain: governed language agents produced higher strategy diversity than hand-coded agents using the same behavioural theory with parameterized heterogeneity, because language-based reasoning generates qualitatively different strategies that threshold-based rules cannot represent. A fuzzy Q-learning baseline (Hung & Yang, 2021) achieved comparable demand ratios but near-zero demand\u2013Mead coupling (r = 0.057 versus 0.547 governed), confirming that the adaptive exploitation pattern requires the natural-language reasoning format, not merely the governance constraints (Supplementary Section S11).',

    'Governance also reveals reasoning patterns that illuminate how agents behave at institutional boundaries. Because the broker records every blocked proposal alongside the agent\u2019s natural-language reasoning, governance creates an audit trail of decisions that agents attempted but the institution prevented. During Tier 3 shortage, for example, an agent reasoned: \u201cmy previous aggressive increase failed... given my scepticism of forecasts, I will cautiously increase\u201d \u2014 acknowledging failure while repeating it (see Supplementary Section S2 for paired reasoning traces). This pre-curtailment acceleration has empirical precedent in prior-appropriation systems where users historically increase diversions in anticipation of curtailment. Our governance architecture suppresses these proposals; future designs could distinguish between physically impossible actions and institutionally unsanctioned-but-observed strategies, preserving them as diagnostic windows into reasoning under scarcity. This distinguishes boundary constraints (what agents cannot do) from prescriptive constraints (what agents should do) \u2014 the majority of our rules enforce the former, leaving strategy selection to the agent.',

    'Several scope conditions apply. Only one model (Gemma-3 4B) was tested in the irrigation domain; the multi-scale analysis is limited to flood. Both model families are instruction-tuned decoder-only transformers; generalization to reasoning-specialized architectures remains untested. The reservoir model simplifies the full CRSS system to a single reservoir; results should be interpreted as stylized institutional dynamics. Each domain required substantial configuration (personas, validators, skill registries); transferability refers to the governance architecture, not the configuration effort. Despite these limitations, the consistency of the governance effect \u2014 positive for all six model scales, across both water domains, and against a hand-coded baseline \u2014 suggests that institutional boundary enforcement is a robust mechanism for enabling adaptive strategy diversity in language-based water simulation. The question for water research shifts from whether language-based agents can represent human water reasoning to what institutional designs best govern that reasoning \u2014 and what those designs reveal about the adaptive capacity institutions create.',
]
for para in discussion_paras:
    add_paragraph(para)

doc.add_page_break()

# =====================================================================
# METHODS
# =====================================================================
doc.add_heading('Methods', level=1)

# Read methods from file and parse sections
import pathlib
methods_path = pathlib.Path(r'C:\Users\wenyu\Desktop\Lehigh\governed_broker_framework\paper\nature_water\drafts\methods_v3.md')
methods_text = methods_path.read_text(encoding='utf-8')

# Extract content between first --- and second ---
parts = methods_text.split('---')
if len(parts) >= 3:
    methods_body = parts[1] if '## Methods' not in parts[1] else parts[1]
    # Find the actual methods content
    for i, part in enumerate(parts):
        if '## Methods' in part:
            methods_body = part
            break

# Parse methods into sections
lines = methods_body.strip().split('\n')
current_heading = None
current_text = []

def flush_section():
    if current_heading and current_text:
        doc.add_heading(current_heading, level=2)
        # Join lines into paragraphs (split on double newlines)
        full_text = '\n'.join(current_text)
        paragraphs = re.split(r'\n\n+', full_text.strip())
        for para_text in paragraphs:
            cleaned = para_text.strip()
            if not cleaned:
                continue
            # Handle numbered lists
            if re.match(r'^\d+\.', cleaned):
                # It's a numbered list block
                items = re.split(r'\n(?=\d+\.)', cleaned)
                for item in items:
                    item = item.strip()
                    if item:
                        # Remove markdown bold
                        item = re.sub(r'\*\*(.+?)\*\*', r'\1', item)
                        add_paragraph(item)
            elif cleaned.startswith('- '):
                # Bullet list
                items = cleaned.split('\n')
                for item in items:
                    item = item.strip().lstrip('- ')
                    item = re.sub(r'\*\*(.+?)\*\*', r'\1', item)
                    if item:
                        p = doc.add_paragraph(style='List Bullet')
                        run = p.add_run(item)
            elif cleaned.startswith('Storage(t+1)'):
                # Equation
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(cleaned)
                run.italic = True
            elif cleaned.startswith('where '):
                add_paragraph(cleaned)
            else:
                cleaned = re.sub(r'\*\*(.+?)\*\*', r'\1', cleaned)
                add_paragraph(cleaned)

for line in lines:
    if line.startswith('### '):
        flush_section()
        current_heading = line.replace('### ', '').strip()
        current_text = []
    elif line.startswith('## Methods'):
        continue  # skip the top-level heading
    else:
        current_text.append(line)

flush_section()

# =====================================================================
# Save
# =====================================================================
output_path = r'C:\Users\wenyu\Desktop\Lehigh\governed_broker_framework\paper\nature_water\drafts\NW_Full_v14.docx'
doc.save(output_path)
print(f'Saved: {output_path}')

# Word count estimate
all_text = abstract_text
for p in intro_paragraphs:
    all_text += ' ' + p
for p in r1_paras + r2_paras + r3_paras + r4_paras:
    all_text += ' ' + p
all_text += ' ' + r4_closing
for p in discussion_paras:
    all_text += ' ' + p

wc = len(all_text.split())
print(f'Main text word count (Abstract + Intro + Results + Discussion): ~{wc}')
