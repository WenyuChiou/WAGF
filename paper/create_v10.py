"""
Create SAGE_WRR_Paper_v10.docx from v9 with Section 5 irrigation edits.

Changes:
1. [52] 5.1 Setup paragraph 1: Add specific experimental parameters
2. [53] 5.1 Setup paragraph 2: Add SI cross-references
3. [57] 5.2 Results paragraph 1: Add quantitative evidence
4. [58] 5.2 Results paragraph 2: Add cold-start explanation
5. [59] 5.2 Results paragraph 3: Add governance metrics + cluster behavior
6. [61] Figure 3 caption: Add SI cross-references
"""
import sys
import copy
from pathlib import Path

sys.stdout.reconfigure(encoding="utf-8")

from docx import Document

PAPER_DIR = Path(__file__).parent
SRC = PAPER_DIR / "SAGE_WRR_Paper_v9.docx"
DST = PAPER_DIR / "SAGE_WRR_Paper_v10.docx"

doc = Document(str(SRC))

# ── Replacement map: paragraph index → new text ──
# We replace full paragraph text while preserving style.

REPLACEMENTS = {

    # [52] 5.1 Setup paragraph 1 — add specifics
    52: (
        "To evaluate transferability, we configure WAGF in a Colorado River "
        "irrigation-demand setting while keeping the core governance runtime "
        "unchanged from the flood case. The experiment simulates 78 CRSS irrigation "
        "districts over 42 years (2019\u20132060) using Gemma 3 4B, with agent personas "
        "derived from k-means clustering of the Fuzzy Q-Learning (FQL) parameters in "
        "Hung and Yang (2021, Table 2): 67 aggressive, 5 forward-looking conservative, "
        "and 6 myopic conservative agents. Domain adaptation is implemented through "
        "configuration-level substitutions: five graduated demand-adjustment skills "
        "(increase_large, increase_small, maintain_demand, decrease_small, decrease_large) "
        "replace the flood action set, 12 custom validators enforce physical and "
        "institutional constraints (Section S6), and a dual-appraisal structure\u2014Water "
        "Scarcity Assessment (WSA) and Adaptive Capacity Assessment (ACA)\u2014replaces "
        "the PMT constructs used in the flood case. Demand-change magnitude is sampled "
        "from persona-scaled Gaussian distributions at execution time rather than "
        "specified by the LLM (Section S7)."
    ),

    # [53] 5.1 Setup paragraph 2 — add SI refs + mass balance
    53: (
        "Governance is applied in ordered layers: identity and physical feasibility "
        "constraints first, followed by construct-conditioned coherence checks, "
        "retry-with-feedback (Section S2), and auditable execution outcomes. "
        "The environment couples agent demand to Lake Mead reservoir dynamics "
        "through an annual mass balance model (Section S8), creating bidirectional "
        "feedback: agent diversions lower storage, triggering shortage tiers that "
        "curtail future diversions and activate governance rules. "
        "Figure 3 shows that this reconfigured setup yields plausible long-horizon "
        "system behavior under institutional and hydrologic constraints, while "
        "preserving bounded behavioral heterogeneity at the agent level. The "
        "irrigation case is therefore presented as transferability evidence: the same "
        "framework operates across water domains without redesigning core architecture."
    ),

    # [57] 5.2 Results paragraph 1 — add quantitative evidence
    57: (
        "Figure 3 presents the transferability demonstration. Panel (a) shows "
        "system-level demand trajectories: steady-state mean demand (Y6\u201342) is "
        "5.87 MAF/yr (1.003\u00d7 the CRSS static baseline), with coefficient of "
        "variation 5.3% and 88% of years falling within the \u00b110% reference "
        "corridor. Panel (b) shows Lake Mead elevation dynamics, ranging from "
        "1,003 to 1,179 ft across the 42-year horizon with 12 shortage years "
        "(Tier 1: 5, Tier 2: 2, Tier 3: 5). These metrics fall within the "
        "range reported by Hung and Yang (2021) for FQL agents under comparable "
        "climate scenarios, achieved through governance constraints rather than "
        "reward-based Q-value convergence."
    ),

    # [58] 5.2 Results paragraph 2 — add cold-start + governance metrics
    58: (
        "The first five years constitute a memory initialization transient "
        "(analogous to spin-up in physical models), during which agents lack "
        "episodic context for informed decision-making: cold-start mean demand "
        "is 4.76 MAF with CoV 11.4%, compared to 6.02 MAF and 5.3% in "
        "steady state. Of 3,276 agent-year decisions, 37.7% are approved on "
        "first attempt, 22.4% succeed after governance retry, and 39.8% are "
        "rejected and fall back to maintain_demand. This 60% intervention rate "
        "reflects a structural feature of bounded-rationality LLM agents under "
        "chronic drought, where governance compensates for the absence of a "
        "reward signal (Section S9)."
    ),

    # [59] 5.2 Results paragraph 3 — add H_norm compression + clusters
    59: (
        "Governance visibly narrows the executable action space while retaining "
        "interpretable differences across behavioral clusters. Normalized Shannon "
        "entropy drops from H_norm = 0.74 (proposed) to 0.39 (executed), a 47% "
        "compression that quantifies institutional constraint strength. Critically, "
        "cluster differentiation is preserved: aggressive agents face 43 "
        "percentage-point governance compression (propose 60% increase \u2192 execute "
        "17%), while myopic agents face near-zero compression (98% maintain in "
        "both proposed and executed). This indicates bounded heterogeneity rather "
        "than homogeneous lock-step behavior\u2014the qualitative behavioral ordering "
        "from FQL k-means clusters is maintained through governance rules rather "
        "than individually calibrated penalty sensitivities (Section S7)."
    ),

    # [61] Figure 3 caption — add SI refs
    61: (
        "Figure 3. Irrigation case study: 78 CRSS districts, 42 years, Gemma 3 4B, "
        "Phase C governance. (a) Annual aggregate water demand. Dashed indigo: CRSS "
        "static baseline (USBR, 2012). Solid teal: WAGF governed request. Dotted "
        "blue: actual diversion after curtailment. Shaded band: \u00b110% CRSS reference "
        "range. The cold-start transient (2019\u20132023) reflects zero-memory "
        "initialization; steady-state demand (2024\u20132060) tracks the CRSS baseline "
        "within the \u00b110% corridor (88% of years). (b) Lake Mead elevation "
        "trajectory with DCP shortage tier bands. Tier thresholds at 1,075 ft "
        "(Tier 1, 5% curtailment), 1,050 ft (Tier 2, 10%), and 1,025 ft (Tier 3, "
        "20%) follow the 2019 Drought Contingency Plan. Elevation ranges from 1,003 "
        "to 1,179 ft across the 42-year horizon, with 12 shortage years (Tier 1: 5, "
        "Tier 2: 2, Tier 3: 5). The bidirectional coupling between panels is the core "
        "mechanism: agent demand decisions (a) affect reservoir storage and elevation "
        "(b), which in turn determines shortage tiers and curtailment ratios fed back "
        "to agents. Complete governance rule specifications, FQL-to-persona cluster "
        "mapping, mass balance equations, and production summary statistics are "
        "provided in Sections S6\u2013S9."
    ),
}

# ── Apply replacements ──
changes = 0
for idx, new_text in REPLACEMENTS.items():
    para = doc.paragraphs[idx]
    old_text = para.text.strip()

    # Clear existing runs
    for run in para.runs:
        run.text = ""

    # Set new text in first run (or add one)
    if para.runs:
        para.runs[0].text = new_text
    else:
        para.add_run(new_text)

    changes += 1
    # Show first 80 chars of old vs new
    print(f"[{idx}] OLD: {old_text[:80]}...")
    print(f"[{idx}] NEW: {new_text[:80]}...")
    print()

doc.save(str(DST))
print(f"Saved: {DST}")
print(f"Total changes: {changes}")
